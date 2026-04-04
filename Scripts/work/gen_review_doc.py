"""Generate the code review DOCX report."""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

title = doc.add_heading('MarkFlow Codebase Review', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Comprehensive Code Review Report')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run(f'Generated: {datetime.datetime.now().strftime("%B %d, %Y")} | Version: v0.19.6.6')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(130, 130, 130)

doc.add_paragraph()

score_p = doc.add_paragraph()
score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = score_p.add_run('Overall Score: 68 / 100')
run_s.bold = True
run_s.font.size = Pt(24)
run_s.font.color.rgb = RGBColor(220, 120, 20)

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'The MarkFlow codebase demonstrates strong architecture and engineering quality for its scope. '
    'The handler registry pattern, domain-split DB package, scan coordinator, and adaptive throttling '
    'are well-designed. Code follows its own conventions consistently with good structlog usage and clear naming.'
)
doc.add_paragraph(
    'The score of 68/100 is pulled down primarily by production-readiness gaps: configuration defaults '
    'that are wide open, missing input validation guardrails, and a migration error swallowing bug. '
    'Most critical/high items are quick config fixes. Two items need real engineering time: migration '
    'error handling and testing security-critical routes.'
)

# Critical Issues
doc.add_heading('Critical Issues (Fix Before Production)', level=1)

t = doc.add_table(rows=4, cols=4)
t.style = 'Medium Shading 1 Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
hdr[0].text = '#'
hdr[1].text = 'Finding'
hdr[2].text = 'File'
hdr[3].text = 'Fix Effort'

critical_data = [
    ('1', 'DEV_BYPASS_AUTH=true hardcoded in docker-compose.yml as a literal, not overridable via .env. Every deployment using this file has no authentication.', 'docker-compose.yml:33', '1 min'),
    ('2', 'Migration errors silently swallowed \u2014 except Exception: pass catches ALL failures (disk full, permissions, schema conflicts), then marks migration as applied. Failed CREATE TABLE becomes permanently "done".', 'core/db/schema.py:638-648', '30 min'),
    ('3', 'SQL injection via counter param in increment_bulk_job_counter \u2014 f-string interpolation of column name. No exploit path today (all callers use literals), but no whitelist guard.', 'core/db/bulk.py:69-76', '10 min'),
]

for i, row_data in enumerate(critical_data):
    row = t.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

# High Severity
doc.add_heading('High Severity Issues', level=1)

t2 = doc.add_table(rows=6, cols=3)
t2.style = 'Medium Shading 1 Accent 1'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = t2.rows[0].cells
hdr2[0].text = '#'
hdr2[1].text = 'Finding'
hdr2[2].text = 'File'

high_data = [
    ('4', 'SSRF in /api/llm-providers/ollama-models \u2014 base_url query param used directly in HTTP request with no validation against internal IPs. With DEV_BYPASS_AUTH=true, any user can probe internal services.', 'api/routes/llm_providers.py:87-106'),
    ('5', 'Unbounded memory leak \u2014 _rate_buckets dict in client_log.py grows indefinitely per unique IP, never cleaned up. Accumulates thousands of empty deques over days/weeks.', 'api/routes/client_log.py:25'),
    ('6', 'SECRET_KEY default is public \u2014 "dev-secret-change-in-prod" encrypts stored LLM API keys. Any deployment that forgets to set it = effectively plaintext API keys.', 'docker-compose.yml:29'),
    ('7', 'Meilisearch in development mode \u2014 auth enforcement disabled even with a master key set. Port 7700 exposed to host. Anyone who can reach it has full index access.', 'docker-compose.yml:84'),
    ('8', '**extra_fields keys interpolated as SQL column names \u2014 systemic pattern in upsert_source_file, update_bulk_file, update_bulk_job_status. All current callers safe, but no whitelist guard.', 'core/db/bulk.py:81-135'),
]

for i, row_data in enumerate(high_data):
    row = t2.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

# Important
doc.add_heading('Important Issues', level=1)

t3 = doc.add_table(rows=5, cols=3)
t3.style = 'Medium Shading 1 Accent 1'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = t3.rows[0].cells
hdr3[0].text = '#'
hdr3[1].text = 'Finding'
hdr3[2].text = 'File'

imp_data = [
    ('9', 'Unauthenticated client-log endpoint \u2014 no require_role() on /api/log/client-event, allows log injection at 50/sec/IP when auth is enabled.', 'api/routes/client_log.py:48'),
    ('10', 'MCP server auth token defaults to empty \u2014 port 8001 exposed, no enforcement code. Any process that can reach it has full MCP tool access.', 'mcp_server/server.py:29'),
    ('11', 'No pinned dependency versions \u2014 requirements.txt uses minimums only. Builds on different days can produce different containers with breaking changes or CVEs.', 'requirements.txt'),
    ('12', 'Batch upsert fallback silently drops files \u2014 failed files logged at WARNING but become phantom "pending" entries that accumulate across runs forever.', 'core/db/bulk.py:288-306'),
]

for i, row_data in enumerate(imp_data):
    row = t3.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()

# Dimension Assessment
doc.add_heading('Assessment by Dimension', level=1)

t4 = doc.add_table(rows=11, cols=2)
t4.style = 'Medium Shading 1 Accent 1'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = t4.rows[0].cells
hdr4[0].text = 'Dimension'
hdr4[1].text = 'Assessment'

dims = [
    ('Correctness', 'Good overall. Scan/convert pipeline well-tested in practice. Migration error swallowing (#2) is the standout risk.'),
    ('Security', 'Multiple hardened-in-dev-but-open-in-practice issues. Auth bypass, dev-mode Meilisearch, empty MCP token, public SECRET_KEY default. No active SQL exploit paths, but no guard rails either.'),
    ('Error Handling', 'Strong "one bad file never crashes a batch" pattern, well-applied in workers. Weak in migrations and batch upsert fallback.'),
    ('Readability', 'Good. Consistent structlog usage, clear naming, well-organized modules. The **extra_fields SQL builder pattern is the main pain point.'),
    ('Architecture', 'Solid. Clean separation: core/db/ domain split, format handler registry, scan coordinator, pipeline layers. main.py lifespan growing long but not a defect.'),
    ('Performance', 'Adaptive scan throttling, batched DB writes, async DB overlap \u2014 all well-done. Only leak is the rate-limiter dict (#5).'),
    ('Test Coverage', 'Tests exist for core paths but NO tests for flag_manager, flags.py, llm_providers.py, or pipeline.py \u2014 all security-critical paths.'),
    ('Input Validation', 'Present but inconsistent. Sort/filter params whitelisted in some routes, not others.'),
    ('Dependencies', 'No lock file, no pinned versions. Known reproducibility risk.'),
    ('Pattern Consistency', 'Strong. New code follows existing conventions (structlog, handler registry, DB helpers). Recent additions follow established patterns cleanly.'),
]

for i, row_data in enumerate(dims):
    row = t4.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()

# Recommendations
doc.add_heading('Recommended Fix Priority', level=1)

doc.add_heading('Immediate (config changes, no code)', level=2)
for item in [
    'Change DEV_BYPASS_AUTH to ${DEV_BYPASS_AUTH:-false} in docker-compose.yml',
    'Set MEILI_ENV=production in docker-compose.yml',
    'Generate a proper SECRET_KEY at install time (openssl rand -hex 32)',
    'Set a non-empty MCP_AUTH_TOKEN default or enforce auth in MCP server',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Short-term (small code changes)', level=2)
for item in [
    'Add column-name whitelist to increment_bulk_job_counter and all **fields update functions',
    'Validate base_url in Ollama models endpoint against private IP ranges',
    'Add require_role() to client-event endpoint',
    'Add TTL cleanup to _rate_buckets dict',
    'Pin dependency versions with pip freeze',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Engineering effort required', level=2)
for item in [
    'Rewrite migration error handling: only catch OperationalError for ALTER TABLE, let CREATE TABLE errors propagate',
    'Add test coverage for flag_manager, flags.py, llm_providers.py, and pipeline.py',
    'Audit all **extra_fields patterns and add schema-column whitelists',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run('Report generated by Claude Code \u2014 MarkFlow Code Review')
run_f.font.size = Pt(8)
run_f.font.color.rgb = RGBColor(150, 150, 150)

doc.save('/app/data/markflow-code-review-2026-04-03.docx')
print('DONE: /app/data/markflow-code-review-2026-04-03.docx')
