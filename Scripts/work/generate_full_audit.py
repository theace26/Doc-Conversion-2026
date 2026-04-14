"""Generate MarkFlow combined audit + spec review + plan .docx report."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def checkbox(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[ ]  {text}")
    run.font.size = Pt(10)
    return p


def bold_line(label, value):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(value)
    return p


def status_badge(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


# ==================================================================
# TITLE PAGE
# ==================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("MarkFlow", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.size = Pt(36)

sub = doc.add_heading("Health Audit, Spec Review & Implementation Plan", level=0)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.color.rgb = RGBColor(0x44, 0x6E, 0x9B)
    run.font.size = Pt(18)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Version 0.22.19  |  Branch: vector  |  Date: 2026-04-09\n").font.size = Pt(11)
meta.add_run("Prepared for Xerxes Shelley\n").font.size = Pt(11)
meta.add_run("\nPart 1: Live System Health Audit (log analysis + DB inspection)\n").font.size = Pt(10)
meta.add_run("Part 2: Specification Review Answers (C1-C6, M1-M5, S1-S6)\n").font.size = Pt(10)
meta.add_run("Part 3: Implementation Plan for Remaining Work").font.size = Pt(10)

doc.add_page_break()

# ==================================================================
# TABLE OF CONTENTS
# ==================================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "",
    "PART 1: LIVE SYSTEM HEALTH AUDIT",
    "  1. Executive Summary",
    "  2. Key Metrics Snapshot",
    "  3. Critical Findings (6 issues)",
    "  4. Deep Pipeline Analysis",
    "  5. Prioritized Recommendations (13 items)",
    "  6. Implementation Tracking Checklist",
    "",
    "PART 2: SPECIFICATION REVIEW ANSWERS",
    "  7. Critical Recommendations (C1-C6)",
    "  8. Medium Recommendations (M1-M5)",
    "  9. Suggested Improvements (S1-S6)",
    "",
    "PART 3: IMPLEMENTATION PLAN",
    "  10. Items Requiring Implementation",
    "  11. Phased Execution Plan",
    "  12. Master Checklist",
    "",
    "APPENDICES",
    "  A. Log Analysis Methodology",
    "  B. Container State",
    "  C. SQLite Configuration",
    "  D. bulk_files Duplication Distribution",
]
for item in toc_items:
    if not item:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.font.size = Pt(10)
        if not item.startswith("  "):
            run.bold = True

doc.add_page_break()

# ██████████████████████████████████████████████████████████████████
# PART 1: HEALTH AUDIT
# ██████████████████████████████████████████████████████████████████
part1 = doc.add_heading("PART 1: LIVE SYSTEM HEALTH AUDIT", level=0)
for run in part1.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# --- 1. Executive Summary ---
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The MarkFlow stack is operational (14-hour uptime, all 4 containers healthy), "
    "scanning is active, and the lifecycle state machine is functioning. However, "
    "log analysis and DB inspection reveal six systemic issues that degrade performance, "
    "inflate database size, and block the image analysis pipeline."
)
doc.add_paragraph(
    "The most critical finding is a polling-driven status endpoint that blocks for up to "
    "7 minutes under load, creating cascading DB contention across the entire system. "
    "Combined with relentless frontend polling (2,200 requests/hour with zero users), "
    "a 100% failure rate on vision analysis (MIME mismatch), and 3.1x row inflation in "
    "bulk_files, the system is doing significant wasted work."
)
doc.add_paragraph(
    "Deep pipeline analysis reveals additional architectural bottlenecks: per-file DB "
    "writes in the worker loop (276K individual UPDATEs), a conversion semaphore limited "
    "to 3 (starving 5 of 8 workers), and file I/O operations held inside DB transactions "
    "causing lock contention."
)

# --- 2. Key Metrics ---
doc.add_heading("2. Key Metrics Snapshot", level=1)

add_table(
    ["Metric", "Value", "Assessment"],
    [
        ["source_files (total)", "88,536", "Healthy"],
        ["  - active", "83,091", ""],
        ["  - marked_for_deletion", "2,467", ""],
        ["  - in_trash", "2,978", ""],
        ["bulk_files (total)", "276,320", "!! 3.1x ratio to source_files"],
        ["  - pending", "263,609", "!! 95% unconverted backlog"],
        ["  - converted", "4,346 (1.6%)", "!! Pipeline barely converting"],
        ["  - unrecognized", "7,748", ""],
        ["  - failed", "177", ""],
        ["  - skipped", "440", ""],
        ["Database size", "312.9 MB (WAL mode)", "Growing fast"],
        ["SQLite busy_timeout", "5,000 ms", "!! Too short"],
        ["lifecycle.trashed events (today)", "35,341", "!! 2,500/hour"],
        ["Vision API success rate", "20% (29/115 fail)", "!! MIME mismatch"],
        ["DB contention log", "2.4M lines", "!! Chronic"],
        ["HTTP requests (status polling)", "~49,000/day", "!! Zero-user baseline"],
    ],
)

doc.add_paragraph()
doc.add_heading("Query Benchmark (direct SQLite, no contention)", level=2)
add_table(
    ["Query", "Result", "Duration"],
    [
        ["Active source_files COUNT", "83,091", "935 ms"],
        ["Pending conversion (NOT EXISTS join)", "82,898", "200 ms"],
        ["Failed (NOT EXISTS join)", "70", "160 ms"],
        ["Unrecognized (NOT EXISTS join)", "1,638", "196 ms"],
        ["Total (all 4 queries)", "-", "~1,500 ms (cold)"],
    ],
)
doc.add_paragraph(
    "Under contention (during active scans), these same queries balloon to 5-431 seconds."
)

doc.add_page_break()

# --- 3. Critical Findings ---
doc.add_heading("3. Critical Findings", level=1)

# 3.1
doc.add_heading("3.1  [CRITICAL] Pipeline Stats Endpoint Blocks Up to 7 Minutes", level=2)
bold_line("Endpoint: ", "/api/pipeline/stats  &  /api/pipeline/status")
bold_line("File: ", "api/routes/pipeline.py, lines 226-298")
doc.add_paragraph(
    "At startup (04:33 UTC), 8 queued requests each took 5-7 minutes (peak: 431s). "
    "Runs four NOT EXISTS subqueries across 276K bulk_files rows plus three Meilisearch "
    "HTTP calls, polled every 5 seconds by the frontend."
)
add_table(
    ["Duration", "Endpoint", "Timestamp (UTC)"],
    [
        ["431,473 ms (7.2 min)", "/api/pipeline/status", "04:33:54"],
        ["430,504 ms (7.2 min)", "/api/pipeline/stats", "04:33:54"],
        ["408,230 ms (6.8 min)", "/api/pipeline/stats", "04:33:54"],
        ["30,350 ms", "/api/pipeline/stats", "04:34:32"],
        ["5,287 ms", "/api/pipeline/stats", "05:23:26"],
    ],
)
bold_line("Root cause: ", "Full table scans on 276K rows, no result caching, polled 838x/hour.")
bold_line("Impact: ", "Holds SQLite read lock, blocking all concurrent writes.")

# 3.2
doc.add_heading("3.2  [HIGH] Frontend Polling Generates ~2,200 Requests/Hour", level=2)
bold_line("File: ", "static/js/global-status-bar.js, line 38")
add_table(
    ["Endpoint", "Requests/Hour", "Effective Interval"],
    [
        ["/api/admin/active-jobs", "1,351", "~2.7 sec"],
        ["/api/pipeline/stats", "838", "~4.3 sec"],
        ["/api/scanner/progress", "838", "~4.3 sec"],
        ["/api/auto-convert/status", "172", "~21 sec"],
    ],
)
doc.add_paragraph("With zero users, this generates 49,000+ DB-backed requests/day.")

# 3.3
doc.add_heading("3.3  [HIGH] Vision Adapter Failing 100% -- MIME Mismatch", level=2)
bold_line("File: ", "core/vision_adapter.py")
doc.add_paragraph(
    "All 115 image analysis batches failed with HTTP 400 from Anthropic API: "
    "files with .jpg extension are actually GIFs. The adapter trusts the extension "
    "instead of sniffing magic bytes. ~1,150 images not getting AI descriptions."
)

# 3.4
doc.add_heading("3.4  [HIGH] bulk_files Row Inflation -- 3.1x source_files", level=2)
doc.add_paragraph(
    "276,320 bulk_files rows for 88,536 source_files. Each bulk scan job creates new "
    "rows keyed by (job_id, source_path), so re-scanning creates duplicates."
)
add_table(
    ["Rows per Path", "Paths", "Wasted Rows"],
    [
        ["1", "23,729", "0"],
        ["2", "27,187", "27,187"],
        ["3", "20,115", "40,230"],
        ["4-6", "3,925", "~13,164"],
        ["7-8", "9,981", "~65,174"],
        ["10-22", "~3,600", "~42,000"],
        ["TOTAL", "88,536", "~187,784 wasted"],
    ],
)

# 3.5
doc.add_heading("3.5  [MEDIUM] Lifecycle Churn -- 35K Trash Events in 14 Hours", level=2)
doc.add_paragraph(
    "With lifecycle_grace_period_hours at 12 (testing; production: 36+) and "
    "lifecycle_trash_retention_days at 7 (production: 60+), files cycle rapidly. "
    "35,341 trash events today at ~2,500/hour. Expected with test timers but generates "
    "massive log volume and DB write pressure."
)

# 3.6
doc.add_heading("3.6  [MEDIUM] Chronic DB Contention -- 2.4M Lines", level=2)
add_table(
    ["Error Source", "Count", "Type"],
    [
        ["scheduler.trash_move_failed", "21", "database is locked"],
        ["analysis_worker.drain_db_locked_skip", "16", "database is locked"],
        ["lifecycle_scan.deletion_detection_error", "12", "database is locked"],
        ["bulk_job_fatal", "5", "fatal error"],
        ["Trash expiry skipped (max instances)", "8", "job overlap"],
        ["Lifecycle scan skipped (max instances)", "3", "job overlap"],
    ],
)

doc.add_page_break()

# --- 4. Deep Pipeline Analysis ---
doc.add_heading("4. Deep Pipeline Analysis", level=1)

doc.add_heading("4.1  Scanning (bulk_scanner.py)", level=2)
doc.add_paragraph(
    "The bulk scanner walks the source share and upserts files into bulk_files. Each scan "
    "job creates rows keyed by (job_id, source_path). Re-scanning in a new job creates "
    "entirely new rows -- this is the root cause of 3.1x inflation."
)
bold_line("Key bottleneck: ", "Per-file DB round-trip in the upsert loop (276K individual upserts).")
bold_line("Fix: ", "Cross-job dedup after each scan, or change key to (source_path) only.")

doc.add_heading("4.2  Worker Pool & Conversion (bulk_worker.py, converter.py)", level=2)
doc.add_paragraph(
    "8 async workers pull pending files, but the conversion semaphore is hardcoded to 3. "
    "5 of 8 workers are blocked on the semaphore at any time. Each file also triggers "
    "3-4 individual DB writes for progress counters."
)
add_table(
    ["Per-File DB Write", "Location"],
    [
        ["Skip counter increment", "bulk_worker.py ~line 576"],
        ["Fail counter increment", "bulk_worker.py ~line 632"],
        ["Convert counter increment", "bulk_worker.py ~line 836"],
        ["Record conversion result", "converter.py ~line 639"],
        ["OCR stats update (conditional)", "converter.py ~line 661"],
    ],
)
bold_line("Fix: ", "Batch counter updates every 50 files. Raise semaphore to 6-8.")

doc.add_heading("4.3  Scheduler (scheduler.py)", level=2)
add_table(
    ["Job", "Interval", "Yields to Bulk?", "Status Today"],
    [
        ["lifecycle_scan", "45 min", "Yes", "Skipped 3x"],
        ["trash_expiry", "1 hour", "Yes", "Skipped 8x"],
        ["db_compaction", "Weekly", "Yes", "OK"],
        ["bulk_files_self_correction", "6 hours", "Yes", "OK"],
        ["_pipeline_watchdog", "1 hour", "No", "Missed 1x"],
        ["_run_deferred_conversions", "15 min", "Yes", "Skipped 1x"],
    ],
)
bold_line("Risk: ", "Orphaned jobs (crash/OOM) block all maintenance indefinitely. No stale detection.")

doc.add_heading("4.4  Lifecycle (lifecycle_manager.py)", level=2)
doc.add_paragraph(
    "move_to_trash() performs file I/O (shutil.move) inside a DB transaction. "
    "When the move takes >5s on network storage, the transaction holds the write lock, "
    "causing 21 'database is locked' errors today."
)
bold_line("Fix: ", "Move file I/O outside the transaction boundary.")

doc.add_heading("4.5  DB Connection (connection.py)", level=2)
doc.add_paragraph(
    "Each DB operation opens a new aiosqlite connection (no pooling). 8 workers x 3-4 "
    "writes/file = 24+ concurrent connection attempts. PRAGMA setup runs per connection."
)
bold_line("Fix: ", "Connection pool (3-5 reusable connections). Read-only connection for analytics.")

doc.add_page_break()

# --- 5. Recommendations ---
doc.add_heading("5. Prioritized Recommendations", level=1)

doc.add_heading("P0 -- Critical (fix this week)", level=2)

doc.add_heading("5.1  Cache Pipeline Stats (15-30s TTL)", level=3)
bold_line("Effort: ", "Small (1-2 hours)  |  Files: api/routes/pipeline.py")
doc.add_paragraph(
    "In-memory cache with 15-30s TTL on pipeline_stats() and pipeline_status(). "
    "Eliminates 95%+ of heavy queries. Single biggest performance win."
)
checkbox("Add TTL cache to pipeline_stats()")
checkbox("Add TTL cache to pipeline_status()")
checkbox("Verify cache invalidates on bulk job start/complete")

doc.add_heading("5.2  Fix Vision Adapter MIME Detection", level=3)
bold_line("Effort: ", "Small (30 min)  |  Files: core/vision_adapter.py")
doc.add_paragraph(
    "Detect actual MIME from file magic bytes instead of trusting extension. "
    "Unblocks ~1,150 stalled images immediately."
)
checkbox("Add magic-byte MIME detection")
checkbox("Fallback to extension if magic bytes fail")
checkbox("Re-queue failed analysis_queue entries")

doc.add_heading("P1 -- High (fix within 2 weeks)", level=2)

doc.add_heading("5.3  Reduce Frontend Polling (5s -> 15-30s)", level=3)
bold_line("Effort: ", "Small (30 min)  |  Files: static/js/global-status-bar.js")
checkbox("Increase interval to 15-30s (visible), 60s (hidden)")
checkbox("Audit all page-specific polling timers")

doc.add_heading("5.4  Deduplicate bulk_files", level=3)
bold_line("Effort: ", "Medium (2-4 hours)  |  Files: core/db/bulk.py, core/bulk_scanner.py")
checkbox("Choose dedup strategy (cross-job cleanup vs key change)")
checkbox("One-time cleanup of ~187K duplicate rows")
checkbox("Benchmark pipeline/stats before and after")

doc.add_heading("5.5  Increase SQLite busy_timeout (5s -> 30s)", level=3)
bold_line("Effort: ", "Trivial (5 min)  |  Files: core/db/connection.py")
checkbox("Increase busy_timeout to 30,000ms")

doc.add_heading("5.6  Batch Worker Counter Updates", level=3)
bold_line("Effort: ", "Small (1-2 hours)  |  Files: core/bulk_worker.py")
checkbox("In-memory counter accumulator, flush every 50 files")
checkbox("Verify SSE progress updates still work")

doc.add_heading("P2 -- Medium (next cycle)", level=2)

p2_items = [
    "5.7  Raise conversion semaphore (3 -> 6-8) -- core/converter.py",
    "5.8  Suppress httpx/httpcore debug logging -- core/logging_config.py",
    "5.9  Connection pooling (3-5 reusable connections) -- core/db/connection.py",
    "5.10  Move file I/O outside DB transactions -- core/lifecycle_manager.py",
    "5.11  Preferences caching (5-min TTL) -- core/scheduler.py",
    "5.12  Stale job detection (heartbeat + 30-min timeout) -- core/scheduler.py",
    "5.13  Vector indexing backpressure (bounded semaphore) -- core/bulk_worker.py",
]
for item in p2_items:
    checkbox(item)

doc.add_page_break()

# --- 6. Tracking Checklist ---
doc.add_heading("6. Implementation Tracking Checklist", level=1)

add_table(
    ["#", "Action", "Priority", "Effort", "Status", "Notes"],
    [
        ["5.1", "Cache pipeline stats", "P0", "Small", "[ ]", "Biggest single win"],
        ["5.2", "Fix vision MIME detection", "P0", "Small", "[ ]", "Unblocks 1,150 images"],
        ["5.3", "Reduce frontend polling", "P1", "Small", "[ ]", "49K fewer req/day"],
        ["5.4", "Deduplicate bulk_files", "P1", "Medium", "[ ]", "Delete ~187K rows"],
        ["5.5", "Increase busy_timeout", "P1", "Trivial", "[ ]", "5-minute fix"],
        ["5.6", "Batch counter updates", "P1", "Small", "[ ]", "98% fewer DB writes"],
        ["5.7", "Raise conversion semaphore", "P2", "Trivial", "[ ]", "Profile first"],
        ["5.8", "Suppress httpx logging", "P2", "Trivial", "[ ]", "40K fewer lines/day"],
        ["5.9", "Connection pooling", "P2", "Medium", "[ ]", ""],
        ["5.10", "I/O outside transactions", "P2", "Medium", "[ ]", ""],
        ["5.11", "Preferences caching", "P2", "Small", "[ ]", ""],
        ["5.12", "Stale job detection", "P2", "Medium", "[ ]", ""],
        ["5.13", "Vector backpressure", "P2", "Small", "[ ]", ""],
        ["--", "Restore prod lifecycle timers", "Pre-prod", "N/A", "[ ]", "ACTION REQUIRED"],
    ],
)

doc.add_page_break()

# ██████████████████████████████████████████████████████████████████
# PART 2: SPECIFICATION REVIEW ANSWERS
# ██████████████████████████████████████████████████████████████████
part2 = doc.add_heading("PART 2: SPECIFICATION REVIEW ANSWERS", level=0)
for run in part2.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph(
    "Each item from the March 23, 2026 Specification Review is addressed below. "
    "Status is verified against the live codebase as of 2026-04-09."
)

# ==================================================================
# 7. Critical Recommendations (C1-C6)
# ==================================================================
doc.add_heading("7. Critical Recommendations (C1-C6)", level=1)

# --- C1 ---
doc.add_heading("C1. Round-Trip Fidelity Tiers", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")

doc.add_paragraph(
    "All three fidelity tiers are implemented and functioning:"
)
add_table(
    ["Tier", "Description", "Implementation"],
    [
        ["Tier 1", "Structure (headings, paragraphs, lists, tables, images)",
         "DocxHandler.ingest() via python-docx. Always applied."],
        ["Tier 2", "Styles restored from sidecar when available",
         "generate_sidecar() in core/metadata.py. SHA-256[:16] content-hash keying."],
        ["Tier 3", "Original file patch (highest fidelity)",
         "_patch_from_original() in docx_handler.py:840-915. 80% hash-match threshold."],
    ],
)

doc.add_paragraph(
    "Content-addressable anchoring is fully implemented via compute_content_hash() in "
    "core/document_model.py:189, which returns SHA-256[:16] of normalized content. "
    "This ensures minor Markdown edits do not invalidate the entire sidecar. "
    "Tier 3 preserves original document structure, styles, themes, and settings "
    "while applying Markdown-driven content changes as patches."
)
bold_line("Action needed: ", "None. All three recommendations are implemented.")

# --- C2 ---
doc.add_heading("C2. Mammoth Markdown Output", level=2)
status_badge("STATUS: NOT APPLICABLE -- SUPERIOR APPROACH IN PLACE")

doc.add_paragraph(
    "MarkFlow does NOT use Mammoth's deprecated Markdown output. The codebase uses "
    "direct python-docx extraction, which is actually superior to the recommended "
    "mammoth -> HTML -> markdownify pipeline:"
)
doc.add_paragraph(
    "- Eliminates intermediate format conversions (no HTML step)\n"
    "- Avoids Mammoth's deprecated Markdown output entirely\n"
    "- Uses the same library (python-docx) for both extraction and reconstruction\n"
    "- Simpler debugging with fewer conversion-loss points"
)
doc.add_paragraph(
    "Note: mammoth and markdownify are still listed in requirements.txt but are never "
    "imported anywhere in the codebase. These are legacy artifacts and can be removed "
    "to reduce image size."
)
bold_line("Action needed: ", "Remove unused mammoth and markdownify from requirements.txt.")

# --- C3 ---
doc.add_heading("C3. WeasyPrint Fallback for Headless Environments", level=2)
status_badge("STATUS: PARTIALLY IMPLEMENTED -- fpdf2 FALLBACK MISSING")

doc.add_paragraph(
    "WeasyPrint is the primary PDF export engine (formats/pdf_handler.py:372). "
    "A startup health check exists (core/health.py:113-128) that detects if WeasyPrint "
    "is importable and logs a warning referencing fpdf2 as fallback. Both packages are "
    "in requirements.txt."
)
doc.add_paragraph(
    "However, the actual fpdf2 fallback code path is NOT implemented in the PDF export "
    "logic. If WeasyPrint fails at runtime, the conversion fails entirely."
)

doc.add_heading("ANSWER: Is this a real concern or scope creep?", level=3)
doc.add_paragraph(
    "This is a REAL concern, not scope creep, for two reasons:\n\n"
    "1. MarkFlow runs in Docker (headless by design). WeasyPrint depends on Cairo, Pango, "
    "and GDK-PixBuf -- system libraries that can break on base image updates. If any of "
    "these libraries fail to load, ALL PDF exports fail silently.\n\n"
    "2. The friend-deployment package (documented in memory) targets machines you do not "
    "control. Those machines may have different base images or library versions.\n\n"
    "The fix is small: add an if/else in the PDF export path that tries WeasyPrint first, "
    "catches ImportError, and falls back to fpdf2 with a quality warning in the output."
)
bold_line("Action needed: ", "Implement fpdf2 fallback in formats/pdf_handler.py export path.")

# --- C4 ---
doc.add_heading("C4. Concurrency Model for Batch Processing", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")

add_table(
    ["Recommendation", "Status", "Implementation"],
    [
        ["asyncio.to_thread() for CPU-bound work", "Done",
         "core/converter.py:8,185,629 and core/bulk_worker.py"],
        ["Max concurrency limit (semaphore)", "Done",
         "converter.py:71 -- Semaphore(3), configurable via MAX_CONCURRENT_CONVERSIONS"],
        ["Batch state in SQLite for crash recovery", "Done",
         "core/db/schema.py:46-56 batch_state table; core/db/conversions.py:134-154 upsert"],
    ],
)
doc.add_paragraph(
    "The concurrency model is correct but the semaphore is set conservatively at 3. "
    "See audit finding 5.7 -- raising to 6-8 would improve throughput by utilizing "
    "all 8 workers."
)
bold_line("Action needed: ", "Raise semaphore default (see audit item 5.7).")

# --- C5 ---
doc.add_heading("C5. PDF OCR Detection Heuristic", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")

doc.add_paragraph(
    "Multi-signal OCR detection in core/ocr.py:32-89 (needs_ocr() function):"
)
add_table(
    ["Signal", "Implementation"],
    [
        ["Text length / page area ratio", "Entropy/stddev blank detection"],
        ["Character encoding consistency", "Entropy + edges combo (photo vs text)"],
        ["Bounding box validation", "OCRWord includes bbox tuple (core/ocr_models.py:27)"],
        ["Confidence threshold", "Configurable via preferences (default 80)"],
    ],
)

doc.add_heading("ANSWER: How does this work in a document repository?", level=3)
doc.add_paragraph(
    "In the context of MarkFlow's document repository:\n\n"
    "1. During bulk scanning, each PDF is analyzed page-by-page. The heuristic runs on "
    "every page to determine if OCR is needed (some PDFs have mixed pages -- some digital, "
    "some scanned).\n\n"
    "2. If ANY page triggers the OCR threshold, the file is flagged for OCR processing. "
    "The confidence score is stored in the analysis_queue and visible in the Pipeline Files UI.\n\n"
    "3. The force-OCR flag is not explicitly named 'force_ocr' but is effectively controlled "
    "via the OCR confidence threshold preference. Setting it to 0 forces OCR on everything. "
    "The 'unattended' mode in OCRConfig auto-processes without user review.\n\n"
    "4. For the repository use case, the last bullet (force-OCR per file) could be useful "
    "as a manual override in the Pipeline Files page -- users could right-click a file and "
    "force re-OCR regardless of the heuristic."
)
bold_line("Action needed: ", "Consider adding per-file force-OCR button in Pipeline Files UI (low priority).")

# --- C6 ---
doc.add_heading("C6. PyMuPDF Over pdfplumber for Performance", level=2)
status_badge("STATUS: PARTIALLY IMPLEMENTED -- AUTO-SWITCHING NOT BUILT")

doc.add_paragraph(
    "A pdf_engine preference exists in core/db/preferences.py:19 (default: 'pymupdf'), "
    "but pdfplumber is actually used as the primary ingestion tool in formats/pdf_handler.py. "
    "The preference setting exists but the switching logic is not implemented."
)
doc.add_paragraph(
    "The recommendation to auto-switch based on table detection is not built. Currently "
    "pdfplumber handles everything, which means table extraction is good but general "
    "text extraction is slower than necessary for non-table PDFs."
)
bold_line(
    "Action needed: ",
    "Implement PyMuPDF as primary with auto-fallback to pdfplumber for table-heavy PDFs. "
    "Detection: attempt PyMuPDF extraction first, check for table indicators "
    "(gridlines, repeated column alignment), switch to pdfplumber if detected.",
)

doc.add_page_break()

# ==================================================================
# 8. Medium Recommendations (M1-M5)
# ==================================================================
doc.add_heading("8. Medium Recommendations (M1-M5)", level=1)

# --- M1 ---
doc.add_heading("M1. Image Handling Pipeline", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
add_table(
    ["Recommendation", "Status", "Implementation"],
    [
        ["Content-hash naming", "Done",
         "core/image_handler.py:52 -- sha256[:12] + '.png'"],
        ["Non-web format conversion (EMF/WMF/TIFF -> PNG)", "Done",
         "image_handler.py:20,62-65 -- _CONVERT_FORMATS set"],
        ["Dimension preservation in sidecar", "Done",
         "image_handler.py:84-91 -- width/height from PIL"],
    ],
)
bold_line("Action needed: ", "None.")

# --- M2 ---
doc.add_heading("M2. File Size Limits / Upload Validation", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
add_table(
    ["Recommendation", "Status", "Implementation"],
    [
        ["Disk space checks", "Done", "core/converter.py:120-125 validate_file_size()"],
        ["Zip bomb detection", "Done",
         "converter.py:128-153 -- ZIP_BOMB_RATIO=200"],
        ["Max upload size", "Done",
         "DEFAULT_MAX_FILE_MB=100 (env: MAX_UPLOAD_MB)"],
        ["API validation", "Done",
         "api/routes/convert.py:54-70"],
    ],
)
bold_line("Action needed: ", "None.")

# --- M3 ---
doc.add_heading("M3. SQLite Concurrent Write Limitations", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
add_table(
    ["Recommendation", "Status", "Implementation"],
    [
        ["WAL mode", "Done", "core/db/connection.py:55"],
        ["busy_timeout", "Done (but too low)",
         "connection.py:56 -- set to 5s (audit recommends 30s)"],
        ["Batch history inserts", "Done",
         "db_write_with_retry() with exponential backoff"],
    ],
)
bold_line("Action needed: ", "Increase busy_timeout to 30s (see audit item 5.5).")

# --- M4 ---
doc.add_heading("M4. Output File Cleanup Strategy", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
add_table(
    ["Recommendation", "Status", "Implementation"],
    [
        ["Configurable retention period", "Done",
         "lifecycle_manager.py:29 -- TRASH_RETENTION_DAYS=60"],
        ["Cleanup task on schedule", "Done",
         "scheduler.py -- trash_expiry runs hourly"],
        ["Disk usage indicator", "Done",
         "GET /api/admin/disk-usage endpoint"],
        ["Manual purge button", "Done",
         "Settings UI configures retention_days"],
    ],
)
bold_line("Action needed: ", "None.")

# --- M5 ---
doc.add_heading("M5. PPTX Chart/SmartArt Support", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
add_table(
    ["Recommendation", "Status", "Implementation"],
    [
        ["Chart extraction", "Done",
         "formats/pptx_handler.py:165-179 -- [Chart] placeholder"],
        ["SmartArt handling", "Done",
         "Logged as warnings per documentation"],
        ["Unsupported element warnings", "Done",
         "pptx_handler.py:178 -- model.warnings.append()"],
        ["OLE object handling", "Done",
         "pptx_handler.py:182-194 -- HTML comments"],
    ],
)
bold_line("Action needed: ", "None. Consider rendering charts to PNG via LibreOffice headless (enhancement).")

doc.add_page_break()

# ==================================================================
# 9. Suggested Improvements (S1-S6)
# ==================================================================
doc.add_heading("9. Suggested Improvements (S1-S6)", level=1)

# --- S1 ---
doc.add_heading("S1. Conversion Preview Before Committing", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
doc.add_paragraph(
    "POST /api/convert/preview endpoint (api/routes/convert.py:161-206) returns: "
    "detected format, estimated page count, OCR likelihood, element counts "
    "(headings, paragraphs, tables, images), and warnings."
)

doc.add_heading("ANSWER: How does this work in a headless environment?", level=3)
doc.add_paragraph(
    "The preview endpoint is an API call, not a GUI operation. In a headless environment "
    "(Docker, CI/CD, scripts), you call POST /api/convert/preview with the file, and it "
    "returns a JSON response with the preview data. No browser or display needed.\n\n"
    "For the web UI (which is the primary interface), the preview appears as a modal dialog "
    "before the user confirms the conversion. In headless/API mode, the caller decides "
    "whether to proceed based on the JSON response.\n\n"
    "This works perfectly in Docker because the FastAPI server handles everything "
    "server-side -- the 'preview' is just lightweight file analysis, not rendering."
)
bold_line("Action needed: ", "None.")

# --- S2 ---
doc.add_heading("S2. Microsoft markitdown as Alternative Engine", level=2)
status_badge("STATUS: INSTALLED BUT NOT INTEGRATED")
doc.add_paragraph(
    "markitdown[all] is listed in requirements.txt:49 as an 'Alternative conversion engine' "
    "but is not wired into the main conversion pipeline."
)

doc.add_heading("ANSWER: Is markitdown a better engine?", level=3)
doc.add_paragraph(
    "No -- markitdown is NOT a better engine for MarkFlow's use case. Here is why:\n\n"
    "1. markitdown is one-directional (Office -> Markdown only). MarkFlow needs bidirectional "
    "conversion (Markdown -> Office too). markitdown cannot reconstruct .docx files.\n\n"
    "2. markitdown does not support fidelity tiers or style sidecars. It produces plain "
    "Markdown with no metadata preservation -- no fonts, colors, spacing, or layout data.\n\n"
    "3. MarkFlow's python-docx direct extraction already gives better control over the "
    "conversion pipeline, with content-hash keying for round-trip fidelity.\n\n"
    "4. markitdown IS useful as a validation reference -- you could compare its output against "
    "MarkFlow's output to catch edge cases. This is a low-priority enhancement.\n\n"
    "Recommendation: Do NOT replace the current engine. Optionally use markitdown as a "
    "comparison tool for testing. Remove from requirements.txt if not actively used "
    "to reduce Docker image size."
)
bold_line("Action needed: ", "Remove from requirements.txt unless used for validation testing.")

# --- S3 ---
doc.add_heading("S3. Format-Agnostic Intermediate Representation", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
doc.add_paragraph(
    "The DocumentModel (core/document_model.py) is the format-agnostic intermediate "
    "representation. All format handlers convert to/from this model, reducing the "
    "converter matrix from NxM to N+M."
)

doc.add_heading("ANSWER: Benefits to MarkFlow as built now", level=3)
doc.add_paragraph(
    "The DocumentModel is already providing these benefits:\n\n"
    "1. Adding a new format requires only one handler (format <-> DocumentModel), not "
    "one handler per target format. MarkFlow supports ~100 extensions with ~15 handlers.\n\n"
    "2. Style metadata is carried natively in the model (ElementType enum, content_hash, "
    "style_data dict), reducing sidecar complexity.\n\n"
    "3. The content_hash on each Element enables sidecar matching independent of format.\n\n"
    "4. Testing is simpler: verify each handler against DocumentModel, not against every "
    "other format.\n\n"
    "This is already one of MarkFlow's strongest architectural decisions. No changes needed."
)
bold_line("Action needed: ", "None.")

# --- S4 ---
doc.add_heading("S4. Checksum Verification for Round-Trip Testing", level=2)
status_badge("STATUS: PARTIALLY IMPLEMENTED -- NEEDS EXPANSION")
doc.add_paragraph(
    "Round-trip tests exist (tests/test_roundtrip.py, tests/test_round_trip.py) with "
    "structural comparisons for heading count and table dimensions. The content_hash field "
    "in DocumentModel enables structural comparison."
)
doc.add_paragraph(
    "Missing: explicit structural hash function that combines all metrics into a single "
    "comparable hash. Currently comparison is done per-metric in individual test assertions."
)
bold_line(
    "Action needed: ",
    "Implement structural_hash() function in DocumentModel that combines: "
    "heading count + text, table dimensions + cell content, image count + dimensions, "
    "list item count + nesting depth. Add to round-trip test suite.",
)

# --- S5 ---
doc.add_heading("S5. Version the Style Sidecar Schema", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
add_table(
    ["Feature", "Implementation"],
    [
        ["schema_version field", "core/metadata.py:22 -- SCHEMA_VERSION = '1.0.0'"],
        ["Supported versions set", "metadata.py:24 -- SUPPORTED_SCHEMA_VERSIONS"],
        ["Version included in output", "generate_sidecar() at metadata.py:107"],
        ["Version validation on load", "load_sidecar() at metadata.py:119-123"],
        ["Style extractor integration", "core/style_extractor.py:6,23,41-42"],
    ],
)

doc.add_heading("ANSWER: Should we migrate to this?", level=3)
doc.add_paragraph(
    "This is already implemented. The schema is at version 1.0.0 with a migration "
    "framework in place. When you add new style properties (e.g., for a new format), "
    "bump to 1.1.0 and add migration logic in load_sidecar() to handle older sidecars. "
    "The SUPPORTED_SCHEMA_VERSIONS set controls which versions are accepted."
)
bold_line("Action needed: ", "None. Migration framework is in place for future schema changes.")

# --- S6 ---
doc.add_heading("S6. Docker / Docker Compose for Easy Setup", level=2)
status_badge("STATUS: FULLY IMPLEMENTED")
add_table(
    ["Feature", "Implementation"],
    [
        ["Dockerfile", "Multi-stage build, ports 8000 (FastAPI) + 8001 (MCP)"],
        ["Dockerfile.base", "System deps (Tesseract, Poppler, LibreOffice, WeasyPrint)"],
        ["docker-compose.yml", "4 services: markflow, markflow-mcp, meilisearch, qdrant"],
        ["Volume mounts", "/app/data, /mnt/source:ro, /mnt/output-repo, drive mounts"],
        ["Environment variables", ".env.example with all configuration options"],
        ["Friend deployment", "7-step setup documented (folder picker, hardware detect)"],
    ],
)
bold_line("Action needed: ", "None.")

doc.add_page_break()

# ██████████████████████████████████████████████████████████████████
# PART 3: IMPLEMENTATION PLAN
# ██████████████████████████████████████████████████████████████████
part3 = doc.add_heading("PART 3: IMPLEMENTATION PLAN", level=0)
for run in part3.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_heading("10. Items Requiring Implementation", level=1)
doc.add_paragraph(
    "From the audit and spec review, these items need implementation work. "
    "Items already fully implemented are excluded."
)

add_table(
    ["#", "Item", "Source", "Priority", "Effort"],
    [
        ["A", "Cache pipeline stats (15-30s TTL)", "Audit 5.1", "P0", "Small"],
        ["B", "Fix vision adapter MIME detection", "Audit 5.2", "P0", "Small"],
        ["C", "Implement fpdf2 PDF export fallback", "Spec C3", "P1", "Small"],
        ["D", "Implement PyMuPDF/pdfplumber auto-switching", "Spec C6", "P1", "Medium"],
        ["E", "Reduce frontend polling frequency", "Audit 5.3", "P1", "Small"],
        ["F", "Deduplicate bulk_files rows", "Audit 5.4", "P1", "Medium"],
        ["G", "Increase SQLite busy_timeout", "Audit 5.5", "P1", "Trivial"],
        ["H", "Batch worker counter updates", "Audit 5.6", "P1", "Small"],
        ["I", "Raise conversion semaphore to 6-8", "Audit 5.7 + C4", "P2", "Trivial"],
        ["J", "Suppress httpx debug logging", "Audit 5.8", "P2", "Trivial"],
        ["K", "Add structural_hash() for round-trip tests", "Spec S4", "P2", "Small"],
        ["L", "Remove unused deps (mammoth, markdownify)", "Spec C2", "P2", "Trivial"],
        ["M", "Connection pooling", "Audit 5.9", "P2", "Medium"],
        ["N", "Move file I/O outside DB transactions", "Audit 5.10", "P2", "Medium"],
        ["O", "Preferences caching (5-min TTL)", "Audit 5.11", "P2", "Small"],
        ["P", "Stale job detection", "Audit 5.12", "P2", "Medium"],
        ["Q", "Vector indexing backpressure", "Audit 5.13", "P2", "Small"],
    ],
)

doc.add_page_break()

# --- 11. Phased Execution Plan ---
doc.add_heading("11. Phased Execution Plan", level=1)

# Phase 1
doc.add_heading("Phase 1: Critical Fixes (this week)", level=2)
doc.add_paragraph("Goal: Eliminate the two biggest performance/correctness issues.")

doc.add_heading("A. Cache Pipeline Stats", level=3)
doc.add_paragraph(
    "Add an async TTL cache to pipeline_stats() and pipeline_status() in "
    "api/routes/pipeline.py. Implementation approach:"
)
doc.add_paragraph(
    "1. Add module-level cache dict: _stats_cache = {'result': None, 'time': 0}\n"
    "2. At the top of pipeline_stats(), check if time.time() - _stats_cache['time'] < 20\n"
    "3. If cache hit, return _stats_cache['result'] immediately\n"
    "4. If cache miss, run the existing query logic, store result, update time\n"
    "5. Use asyncio.Lock to prevent thundering herd (multiple requests triggering refresh)\n"
    "6. Apply same pattern to pipeline_status()"
)
doc.add_paragraph(
    "Expected impact: Response time drops from 1.5s-431s to <10ms on cache hit. "
    "DB read load cut by ~50%."
)

doc.add_heading("B. Fix Vision Adapter MIME Detection", level=3)
doc.add_paragraph(
    "In core/vision_adapter.py, before constructing the Anthropic API request, "
    "detect the actual MIME type from file magic bytes:"
)
doc.add_paragraph(
    "1. Read the first 32 bytes of the image file\n"
    "2. Check magic bytes: GIF89a/GIF87a = image/gif, FFD8FF = image/jpeg, "
    "89504E47 = image/png, 52494646...57454250 = image/webp\n"
    "3. Use detected MIME instead of extension-based MIME\n"
    "4. Fallback to extension-based if magic bytes are unrecognized\n"
    "5. After deploying, re-queue the 115 failed analysis_queue entries for retry"
)
doc.add_paragraph(
    "Expected impact: Unblocks ~1,150 images for AI description. "
    "Vision success rate goes from 20% to ~95%+."
)

# Phase 2
doc.add_heading("Phase 2: High-Impact Fixes (week 2)", level=2)
doc.add_paragraph("Goal: Reduce DB pressure and fix data inflation.")

doc.add_heading("C. fpdf2 PDF Export Fallback", level=3)
doc.add_paragraph(
    "In formats/pdf_handler.py export method:\n"
    "1. Wrap WeasyPrint call in try/except ImportError\n"
    "2. On failure, fall back to fpdf2 with equivalent output\n"
    "3. Add a warning to the conversion result noting reduced quality\n"
    "4. Log the fallback event for monitoring"
)

doc.add_heading("D. PyMuPDF/pdfplumber Auto-Switching", level=3)
doc.add_paragraph(
    "In formats/pdf_handler.py ingest method:\n"
    "1. First pass: open with PyMuPDF (fitz), extract text\n"
    "2. Table detection: check for gridline patterns or >3 columns of aligned text\n"
    "3. If tables detected: re-extract with pdfplumber for that page\n"
    "4. Merge results: PyMuPDF text + pdfplumber tables\n"
    "5. Wire the pdf_engine preference to allow manual override"
)

doc.add_heading("E. Reduce Frontend Polling", level=3)
doc.add_paragraph(
    "In static/js/global-status-bar.js line 38:\n"
    "1. Change visible interval from 5000 to 20000 (20 seconds)\n"
    "2. Change hidden interval from 30000 to 60000 (60 seconds)\n"
    "3. Audit pipeline-files.html, bulk.html for additional setInterval calls\n"
    "4. Align all polling to 15-30s minimum"
)

doc.add_heading("F. Deduplicate bulk_files", level=3)
doc.add_paragraph(
    "One-time cleanup + ongoing prevention:\n"
    "1. Write a migration that keeps only the LATEST row per source_path "
    "(highest job_id or most recent created_at)\n"
    "2. DELETE the ~187K older duplicate rows\n"
    "3. Run VACUUM to reclaim space (~200MB -> ~120MB)\n"
    "4. Add post-scan cleanup in bulk_scanner.py to prune older duplicates after each scan\n"
    "5. Benchmark pipeline/stats before and after"
)

doc.add_heading("G. Increase busy_timeout", level=3)
doc.add_paragraph(
    "In core/db/connection.py, change PRAGMA busy_timeout from 5000 to 30000. "
    "Single-line change. Deploy and monitor db-contention.log for reduction in "
    "'database is locked' errors."
)

doc.add_heading("H. Batch Worker Counter Updates", level=3)
doc.add_paragraph(
    "In core/bulk_worker.py:\n"
    "1. Add a CounterAccumulator class with dict of pending increments\n"
    "2. On each file completion, increment in-memory counter\n"
    "3. Flush to DB every 50 files or every 5 seconds (whichever first)\n"
    "4. Flush on job completion\n"
    "5. Verify SSE progress events still fire at reasonable intervals"
)

# Phase 3
doc.add_heading("Phase 3: Optimization (weeks 3-4)", level=2)
doc.add_paragraph("Goal: Throughput improvements and code hygiene.")

phase3_items = [
    ("I. Raise Conversion Semaphore",
     "Change MAX_CONCURRENT_CONVERSIONS default from 3 to 6 in core/converter.py:71. "
     "Monitor CPU/memory during conversion to verify the system can sustain it. "
     "With the i7-10750H (6 cores, 12 threads), 6 concurrent conversions should be safe."),
    ("J. Suppress httpx Debug Logging",
     "In core/logging_config.py configure_logging(), add:\n"
     "logging.getLogger('httpcore').setLevel(logging.WARNING)\n"
     "logging.getLogger('httpx').setLevel(logging.WARNING)\n"
     "Same pattern as the existing pdfminer suppression."),
    ("K. Structural Hash for Round-Trip Tests",
     "Add structural_hash() to DocumentModel that combines: heading count + text hashes, "
     "table dimensions + cell content hashes, image count, list item count + nesting. "
     "Return a single SHA-256 for comparison. Add to test_roundtrip.py."),
    ("L. Remove Unused Dependencies",
     "Remove mammoth and markdownify from requirements.txt. Verify no imports exist. "
     "Reduces Docker image size and pip install time."),
    ("M. Connection Pooling",
     "Replace per-call aiosqlite.connect() in core/db/connection.py with a pool of "
     "3-5 reusable connections. Add a dedicated read-only connection for analytics "
     "endpoints (pipeline/stats, scanner/progress)."),
    ("N. Move File I/O Outside DB Transactions",
     "In core/lifecycle_manager.py move_to_trash(): read metadata inside transaction, "
     "commit, perform shutil.move() outside transaction, then update status in a second "
     "quick transaction."),
    ("O. Preferences Caching",
     "Add a module-level cache for frequently-read preferences with a 5-minute TTL. "
     "Invalidate on PUT /api/preferences/<key>."),
    ("P. Stale Job Detection",
     "Add a last_heartbeat column to bulk_jobs. Update heartbeat every 60s during scan. "
     "On startup, mark jobs with status='running' and last_heartbeat > 30 min as 'interrupted'."),
    ("Q. Vector Indexing Backpressure",
     "In core/bulk_worker.py, replace unbounded asyncio.create_task() for Qdrant indexing "
     "with a bounded semaphore (20). If semaphore is full, skip indexing for that file "
     "(it will be picked up on next lifecycle scan)."),
]
for title_text, desc in phase3_items:
    doc.add_heading(title_text, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# --- 12. Master Checklist ---
doc.add_heading("12. Master Checklist", level=1)
doc.add_paragraph(
    "Combined tracking list for all work items from both the audit and spec review. "
    "Print this page and check off items as completed."
)

add_table(
    ["Phase", "#", "Item", "Source", "Effort", "Done"],
    [
        ["1", "A", "Cache pipeline stats (15-30s TTL)", "Audit", "Small", "[ ]"],
        ["1", "B", "Fix vision MIME detection (magic bytes)", "Audit", "Small", "[ ]"],
        ["2", "C", "fpdf2 PDF export fallback", "Spec C3", "Small", "[ ]"],
        ["2", "D", "PyMuPDF/pdfplumber auto-switching", "Spec C6", "Medium", "[ ]"],
        ["2", "E", "Reduce frontend polling (5s -> 20s)", "Audit", "Small", "[ ]"],
        ["2", "F", "Deduplicate bulk_files (~187K rows)", "Audit", "Medium", "[ ]"],
        ["2", "G", "Increase busy_timeout (5s -> 30s)", "Audit", "Trivial", "[ ]"],
        ["2", "H", "Batch worker counter updates", "Audit", "Small", "[ ]"],
        ["3", "I", "Raise conversion semaphore (3 -> 6)", "Audit+C4", "Trivial", "[ ]"],
        ["3", "J", "Suppress httpx debug logging", "Audit", "Trivial", "[ ]"],
        ["3", "K", "Structural hash for round-trip tests", "Spec S4", "Small", "[ ]"],
        ["3", "L", "Remove unused deps (mammoth, etc.)", "Spec C2", "Trivial", "[ ]"],
        ["3", "M", "Connection pooling (3-5 connections)", "Audit", "Medium", "[ ]"],
        ["3", "N", "File I/O outside DB transactions", "Audit", "Medium", "[ ]"],
        ["3", "O", "Preferences caching (5-min TTL)", "Audit", "Small", "[ ]"],
        ["3", "P", "Stale job detection (heartbeat)", "Audit", "Medium", "[ ]"],
        ["3", "Q", "Vector indexing backpressure", "Audit", "Small", "[ ]"],
        ["--", "--", "Restore production lifecycle timers", "Pre-prod", "N/A", "[ ]"],
        ["--", "--", "Address 62 security audit findings", "Pre-prod", "Large", "[ ]"],
    ],
)

doc.add_page_break()

# ==================================================================
# APPENDICES
# ==================================================================
doc.add_heading("Appendix A: Log Analysis Methodology", level=1)
doc.add_paragraph(
    "Analysis performed on 2026-04-09 against live container logs and on-disk log files. "
    "Log window: 04:28-18:13 UTC (~14 hours). All JSON-structured logs parsed and "
    "aggregated by event type and level. HTTP request durations extracted from "
    "request_complete events. DB state queried via python3 inside the running container."
)
add_table(
    ["Source", "Lines Analyzed"],
    [
        ["logs/markflow.log", "150,731"],
        ["logs/markflow-debug.log", "121,445"],
        ["logs/db-contention.log", "2,441,236"],
        ["logs/db-queries.log", "876,645"],
        ["logs/db-active.log", "42,447"],
        ["Total", "3,632,504"],
    ],
)

doc.add_heading("Appendix B: Container State at Time of Audit", level=1)
add_table(
    ["Container", "Image", "Status", "Ports"],
    [
        ["markflow", "doc-conversion-2026-markflow", "Up 14 hours", "8000"],
        ["markflow-mcp", "doc-conversion-2026-markflow-mcp", "Up 14 hours", "8001"],
        ["meilisearch", "getmeili/meilisearch:latest", "Up 14 hours (healthy)", "7700"],
        ["qdrant", "qdrant/qdrant:latest", "Up 14 hours (healthy)", "6333-6334"],
    ],
)

doc.add_heading("Appendix C: SQLite Configuration", level=1)
add_table(
    ["Setting", "Current", "Recommended"],
    [
        ["Journal mode", "WAL", "WAL (no change)"],
        ["WAL autocheckpoint", "1,000 pages", "1,000 (no change)"],
        ["busy_timeout", "5,000 ms", "30,000 ms"],
        ["Database size", "312.9 MB", "~120 MB after dedup"],
        ["Index count", "69", "Add 1-2 composite"],
    ],
)

doc.add_heading("Appendix D: bulk_files Duplication Distribution", level=1)
doc.add_paragraph(
    "How many bulk_files rows exist per unique source_path. Ideally each path = 1 row."
)
add_table(
    ["Rows/Path", "Paths", "Total Rows", "Waste"],
    [
        ["1", "23,729", "23,729", "0"],
        ["2", "27,187", "54,374", "27,187"],
        ["3", "20,115", "60,345", "40,230"],
        ["4-6", "3,925", "17,089", "13,164"],
        ["7-8", "9,981", "75,155", "65,174"],
        ["10-22", "~3,600", "~45,628", "~42,028"],
        ["TOTAL", "88,536", "276,320", "~187,784"],
    ],
)

# Save
out = os.path.join(os.getcwd(), "MarkFlow-Audit-SpecReview-Plan-2026-04-09.docx")
doc.save(out)
print(f"Saved: {out}")
print(f"Size: {os.path.getsize(out) / 1024:.0f} KB")
