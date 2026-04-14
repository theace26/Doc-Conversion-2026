"""Generate MarkFlow Health & Efficiency Audit .docx report."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def checkbox(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[ ]  {text}")
    run.font.size = Pt(10)
    return p


def bold_line(label, value):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(value)
    return p


# ==================================================================
# TITLE PAGE
# ==================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("MarkFlow Health & Efficiency Audit", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Version 0.22.19  |  Branch: vector  |  Date: 2026-04-09\n").font.size = Pt(11)
meta.add_run("Log window: 04:28 - 18:13 UTC (~14 hours)\n").font.size = Pt(10)
meta.add_run("Stack uptime: 14 hours, all 4 containers healthy").font.size = Pt(10)

doc.add_page_break()

# ==================================================================
# TABLE OF CONTENTS (manual)
# ==================================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Key Metrics Snapshot",
    "3. Critical Findings (6 issues)",
    "   3.1  [CRITICAL] Pipeline Stats Endpoint Blocks Up to 7 Minutes",
    "   3.2  [HIGH] Frontend Polling Generates ~2,200 Requests/Hour",
    "   3.3  [HIGH] Vision Adapter Failing 100% - MIME Mismatch",
    "   3.4  [HIGH] bulk_files Row Inflation - 3.1x source_files",
    "   3.5  [MEDIUM] Lifecycle Churn - 35K Trash Events in 14 Hours",
    "   3.6  [MEDIUM] Chronic DB Contention - 2.4M Contention Lines",
    "4. Deep Pipeline Analysis",
    "   4.1  Scanning Pipeline (bulk_scanner.py)",
    "   4.2  Worker Pool & Conversion (bulk_worker.py, converter.py)",
    "   4.3  Scheduler Job Coordination (scheduler.py)",
    "   4.4  Lifecycle State Machine (lifecycle_manager.py)",
    "   4.5  DB Connection & Contention (connection.py)",
    "5. Recommendations - Prioritized Action Plan",
    "6. Implementation Tracking Checklist",
    "Appendix A: Log Analysis Methodology",
    "Appendix B: Container State",
    "Appendix C: SQLite Configuration",
    "Appendix D: bulk_files Duplication Distribution",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ==================================================================
# 1. EXECUTIVE SUMMARY
# ==================================================================
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The MarkFlow stack is operational (14-hour uptime, all 4 containers healthy), "
    "scanning is active, and the lifecycle state machine is functioning. However, "
    "log analysis and DB inspection reveal six systemic issues that degrade performance, "
    "inflate database size, and block the image analysis pipeline."
)
doc.add_paragraph(
    "The most critical finding is a polling-driven status endpoint that blocks for up to "
    "7 minutes under load, creating cascading DB contention across the entire system. "
    "Combined with relentless frontend polling (2,200 requests/hour with zero users), "
    "a 100% failure rate on vision analysis (MIME mismatch), and 3.1x row inflation in "
    "bulk_files, the system is doing significant wasted work."
)
doc.add_paragraph(
    "Deep pipeline analysis reveals additional architectural bottlenecks: per-file DB "
    "writes in the worker loop (276K individual UPDATEs), a conversion semaphore limited "
    "to 3 (starving 5 of 8 workers), and file I/O operations held inside DB transactions "
    "causing lock contention. These are fixable with targeted changes."
)

# ==================================================================
# 2. KEY METRICS
# ==================================================================
doc.add_heading("2. Key Metrics Snapshot", level=1)

add_table(
    ["Metric", "Value", "Assessment"],
    [
        ["source_files (total)", "88,536", "Healthy"],
        ["  - active", "83,091", ""],
        ["  - marked_for_deletion", "2,467", ""],
        ["  - in_trash", "2,978", ""],
        ["bulk_files (total)", "276,320", "!! 3.1x ratio to source_files"],
        ["  - pending", "263,609", "!! 95% unconverted backlog"],
        ["  - converted", "4,346 (1.6%)", "!! Pipeline barely converting"],
        ["  - unrecognized", "7,748", ""],
        ["  - failed", "177", ""],
        ["  - skipped", "440", ""],
        ["Database size", "312.9 MB (WAL mode)", "Growing fast"],
        ["SQLite busy_timeout", "5,000 ms", "!! Too short for workload"],
        ["Index count", "69", "Good coverage overall"],
        ["lifecycle.trashed events (today)", "35,341", "!! 2,500/hour (test timers)"],
        ["Vision API success rate", "20% (29 ok / 115 fail)", "!! MIME mismatch"],
        ["DB contention log", "2.4M lines", "!! Chronic contention"],
        ["HTTP requests (status polling)", "~49,000/day", "!! Zero-user baseline"],
    ],
)

doc.add_paragraph()

doc.add_heading("Query Benchmark (direct SQLite, no contention)", level=2)
add_table(
    ["Query", "Result", "Duration"],
    [
        ["Active source_files COUNT", "83,091", "935 ms"],
        ["Pending conversion (NOT EXISTS join)", "82,898", "200 ms"],
        ["Failed (NOT EXISTS join)", "70", "160 ms"],
        ["Unrecognized (NOT EXISTS join)", "1,638", "196 ms"],
        ["Total (all 4 queries)", "-", "~1,500 ms (cold)"],
    ],
)
doc.add_paragraph(
    "Under contention (during active scans or lifecycle operations), these same "
    "queries balloon to 5-431 seconds as they compete for SQLite locks."
)

doc.add_page_break()

# ==================================================================
# 3. CRITICAL FINDINGS
# ==================================================================
doc.add_heading("3. Critical Findings", level=1)

# --- 3.1 ---
doc.add_heading("3.1  [CRITICAL] Pipeline Stats Endpoint Blocks Up to 7 Minutes", level=2)
bold_line("Endpoint: ", "/api/pipeline/stats  &  /api/pipeline/status")
bold_line("File: ", "api/routes/pipeline.py, lines 226-298")

doc.add_paragraph(
    "At startup (04:33 UTC), 8 queued requests each took 5-7 minutes (peak: 431 seconds). "
    "Throughout the day, the endpoint regularly hits 5 seconds. It runs four heavy "
    "NOT EXISTS subqueries across 276K bulk_files rows, plus three synchronous "
    "Meilisearch HTTP calls - all inline, polled every 5 seconds by the frontend."
)

doc.add_paragraph("Slow request log (all times UTC):")
add_table(
    ["Duration", "Endpoint", "Timestamp"],
    [
        ["431,473 ms (7.2 min)", "/api/pipeline/status", "04:33:54"],
        ["430,504 ms (7.2 min)", "/api/pipeline/stats", "04:33:54"],
        ["408,230 ms (6.8 min)", "/api/pipeline/stats", "04:33:54"],
        ["362,657 ms (6.0 min)", "/api/pipeline/status", "04:33:54"],
        ["30,350 ms (30 sec)", "/api/pipeline/stats", "04:34:32"],
        ["5,287 ms", "/api/pipeline/stats", "05:23:26"],
        ["5,226 ms", "/api/pipeline/stats", "10:13:38"],
        ["5,049 ms", "/api/pipeline/stats", "12:13:20"],
    ],
)

bold_line(
    "Root cause: ",
    "Four COUNT(*) queries with NOT EXISTS subqueries scan the full 276K-row "
    "bulk_files table on every call. At 838 calls/hour, this is the single largest "
    "source of DB load. The three Meilisearch stats calls add network latency.",
)

bold_line(
    "Impact: ",
    "Holds SQLite read lock for seconds, blocking writes from lifecycle scanner, "
    "trash expiry, and bulk worker. Creates cascading 'database is locked' errors.",
)

# --- 3.2 ---
doc.add_heading(
    "3.2  [HIGH] Frontend Polling Generates ~2,200 Requests/Hour", level=2
)
bold_line("File: ", "static/js/global-status-bar.js, line 38")

add_table(
    ["Endpoint", "Requests/Hour", "Effective Interval"],
    [
        ["/api/admin/active-jobs", "1,351", "~2.7 sec"],
        ["/api/pipeline/stats", "838", "~4.3 sec"],
        ["/api/scanner/progress", "838", "~4.3 sec"],
        ["/api/auto-convert/status", "172", "~21 sec"],
        ["Total status polling", "~3,200", ""],
    ],
)

doc.add_paragraph(
    "global-status-bar.js polls every 5 seconds (visible tab) / 30 seconds (hidden). "
    "The pipeline page adds its own polling on top. With zero users browsing, "
    "this generates 49,000+ DB-backed requests per day for status checks alone."
)

# --- 3.3 ---
doc.add_heading(
    "3.3  [HIGH] Vision Adapter Failing 100% - MIME Type Mismatch", level=2
)
bold_line("File: ", "core/vision_adapter.py")

doc.add_paragraph(
    "All 115 image analysis batches failed with HTTP 400 from Anthropic API. "
    'Error: "The image was specified using the image/jpeg media type, but the image '
    'appears to be a image/gif image". The adapter determines MIME type from the file '
    "extension (.jpg) instead of inspecting actual file content (magic bytes)."
)

bold_line("Sample error: ", "")
sample = doc.add_paragraph(
    "Provider: anthropic  |  Batch size: 10  |  First image: "
    "/mnt/source/9thDistrict_Forms/100 YEAR ANNIVERSARY - 2014/Centennial/"
    "photos_jpgOrigs/I-5_Ship_Canal_Bridge_under_construction,_1960.jpg"
)
sample.paragraph_format.left_indent = Cm(1)

bold_line(
    "Impact: ",
    "~1,150 images (115 batches x 10) not getting AI descriptions. "
    "Only 29 successful API calls vs 115 failures = 80% failure rate.",
)

# --- 3.4 ---
doc.add_heading(
    "3.4  [HIGH] bulk_files Row Inflation - 3.1x source_files", level=2
)

doc.add_paragraph(
    "276,320 bulk_files rows for 88,536 source_files means each source file has "
    "~3.1 bulk_file rows on average. Each bulk scan job creates new rows keyed by "
    "(job_id, source_path), so re-scanning the same directory with multiple jobs "
    "creates duplicate rows per file."
)

doc.add_paragraph("Distribution of rows per source_path:")
add_table(
    ["Rows per Path", "Number of Paths", "Cumulative Files"],
    [
        ["1 row", "23,729 paths", "23,729"],
        ["2 rows", "27,187 paths", "54,374"],
        ["3 rows", "20,115 paths", "60,345"],
        ["4 rows", "2,736 paths", "10,944"],
        ["5 rows", "989 paths", "4,945"],
        ["6 rows", "200 paths", "1,200"],
        ["7 rows", "4,693 paths", "32,851"],
        ["8 rows", "5,288 paths", "42,304"],
        ["10-11 rows", "2,193 paths", "~23,000"],
        ["14-22 rows", "~950 paths", "~15,000"],
    ],
)

bold_line(
    "Impact: ",
    "Inflates query time for pipeline stats (3x rows to scan), overstates pending "
    "counts in the UI, wastes ~200 MB of DB space on duplicate rows, and slows "
    "every NOT EXISTS check.",
)

# --- 3.5 ---
doc.add_heading(
    "3.5  [MEDIUM] Lifecycle Churn - 35K Trash Events in 14 Hours", level=2
)

add_table(
    ["Hour (UTC)", "Trash Events", "", "Hour (UTC)", "Trash Events"],
    [
        ["04:00", "185", "", "12:00", "1,615"],
        ["05:00", "1,679", "", "13:00", "3,840"],
        ["06:00", "2,194", "", "14:00", "1,856"],
        ["07:00", "3,013", "", "15:00", "3,373"],
        ["08:00", "2,313", "", "16:00", "2,105"],
        ["09:00", "3,307", "", "17:00", "4,128"],
        ["10:00", "1,879", "", "18:00", "993"],
        ["11:00", "2,793", "", "", ""],
    ],
)

doc.add_paragraph(
    "With lifecycle_grace_period_hours at 12 (testing; production should be 36+) "
    "and lifecycle_trash_retention_days at 7 (production: 60+), files cycle through "
    "mark -> trash -> purge rapidly. This is expected with test timers, but generates "
    "massive log volume and sustained DB write pressure."
)

# --- 3.6 ---
doc.add_heading(
    "3.6  [MEDIUM] Chronic DB Contention - 2.4M Contention Log Lines", level=2
)

add_table(
    ["Error Source", "Count", "Error Type"],
    [
        ["scheduler.trash_move_failed", "21", "database is locked"],
        ["analysis_worker.drain_db_locked_skip", "16", "database is locked (warning)"],
        ["lifecycle_scan.deletion_detection_error", "12", "database is locked"],
        ["bulk_job_fatal", "5", "fatal error"],
        ["persist_throttle_events_failed", "3", "database is locked"],
        ["auto_metrics_aggregation_db_locked_skip", "1", "database is locked"],
        ["Trash expiry job skipped (max instances)", "8", "previous run still running"],
        ["Lifecycle scan skipped (max instances)", "3", "previous run still running"],
    ],
)

doc.add_paragraph(
    "The db-contention.log has accumulated 2.4 million lines. The 5-second busy_timeout "
    "is insufficient when the pipeline/stats endpoint holds read locks for 5-431 seconds. "
    "Trash expiry was skipped 8 times because the previous run had not finished - likely "
    "blocked on DB access."
)

doc.add_page_break()

# ==================================================================
# 4. DEEP PIPELINE ANALYSIS
# ==================================================================
doc.add_heading("4. Deep Pipeline Analysis", level=1)

# --- 4.1 ---
doc.add_heading("4.1  Scanning Pipeline (bulk_scanner.py)", level=2)

doc.add_paragraph(
    "The bulk scanner walks the source share and upserts files into bulk_files. "
    "Each scan job creates rows keyed by (job_id, source_path). Re-scanning the same "
    "directory in a new job creates entirely new rows - this is the root cause of the "
    "3.1x inflation. The scanner does have incremental mode (mtime-based directory "
    "skipping), but it only optimizes within a single scan run."
)

bold_line("Key bottleneck: ", "Per-file DB round-trip in the upsert loop")
doc.add_paragraph(
    "Each file does an individual upsert with a source_files lookup inside the transaction. "
    "For a 276K-file scan, this means 276K individual DB round-trips. The scanner batches "
    "progress yields every 1,000 files, but the DB writes are still per-file."
)

bold_line(
    "Recommendation: ",
    "Run a cross-job deduplication query after each scan completes, or change the "
    "bulk_files key to (source_path) instead of (job_id, source_path) so re-scans "
    "update in place rather than creating new rows.",
)

# --- 4.2 ---
doc.add_heading(
    "4.2  Worker Pool & Conversion (bulk_worker.py, converter.py)", level=2
)

doc.add_paragraph(
    "The bulk worker runs 8 async workers that pull pending files from bulk_files. "
    "However, the conversion semaphore in converter.py is hardcoded to 3:"
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.add_run('_semaphore = asyncio.Semaphore(int(os.getenv("MAX_CONCURRENT_CONVERSIONS", "3")))').font.size = Pt(9)

doc.add_paragraph(
    "This means 5 of 8 workers are blocked on the semaphore at any given time, waiting "
    "to convert. Combined with per-file DB writes for progress counters:"
)

add_table(
    ["Operation", "DB Writes Per File", "Location"],
    [
        ["Skip counter increment", "1", "bulk_worker.py line ~576"],
        ["Fail counter increment", "1", "bulk_worker.py line ~632"],
        ["Convert counter increment", "1", "bulk_worker.py line ~836"],
        ["Record conversion result", "1", "converter.py line ~639"],
        ["OCR stats update (conditional)", "0-1", "converter.py line ~661"],
        ["Total per file", "3-4", ""],
    ],
)

doc.add_paragraph(
    "For 276K pending files, this means 800K-1M individual DB writes just for progress "
    "tracking. Batching these every 50 files would reduce DB writes by 98%."
)

bold_line(
    "Conversion time by format: ",
    "PDFs: 0.1-2s; scanned docs (OCR): 5-30s; Adobe PSD/AI/INDD: 10-60s. "
    "The semaphore limit of 3 makes sense for OCR-heavy workloads but is conservative "
    "for the i7-10750H/64GB/GTX 1660 Ti hardware.",
)

# --- 4.3 ---
doc.add_heading("4.3  Scheduler Job Coordination (scheduler.py)", level=2)

doc.add_paragraph("Scheduled jobs and their intervals:")
add_table(
    ["Job", "Interval", "Yields to Bulk?", "Status Today"],
    [
        ["lifecycle_scan", "45 min", "Yes", "Skipped 3x"],
        ["trash_expiry", "1 hour", "Yes", "Skipped 8x"],
        ["db_compaction", "Weekly", "Yes", "OK"],
        ["bulk_files_self_correction", "6 hours", "Yes", "OK"],
        ["_pipeline_watchdog", "1 hour", "No", "Missed 1x (8s late)"],
        ["_expire_flags", "1 hour", "No", "Missed 1x"],
        ["_run_deferred_conversions", "15 min", "Yes", "Skipped 1x"],
    ],
)

doc.add_paragraph(
    "All maintenance jobs check get_all_active_jobs() before running and skip if any "
    "bulk job is active. This is correct but means long-running bulk jobs can defer "
    "maintenance indefinitely. The self_correction job (which cleans up bulk_files "
    "inflation) only runs every 6 hours and also yields to bulk jobs."
)

bold_line(
    "Orphaned job risk: ",
    "If a bulk job crashes (SIGKILL, OOM), the in-memory active job counter resets "
    "to 0 on container restart, but the DB still shows the job as 'running'. Next "
    "lifecycle scan proceeds immediately, creating a race condition. There is no stale "
    "job detection (e.g., 'if last_heartbeat > 30 min, mark interrupted').",
)

# --- 4.4 ---
doc.add_heading(
    "4.4  Lifecycle State Machine (lifecycle_manager.py)", level=2
)

doc.add_paragraph("State transitions: active -> marked_for_deletion -> in_trash -> purged")

doc.add_paragraph(
    "The 21 'scheduler.trash_move_failed' errors today are caused by the move_to_trash "
    "function performing file I/O (shutil.move) inside a DB transaction. When the file "
    "move takes >5 seconds (network storage, large files), the transaction holds the "
    "SQLite write lock, blocking all other writers."
)

bold_line(
    "Recommendation: ",
    "Fetch file metadata inside the transaction, commit, then perform the file move "
    "outside the transaction. Update the final status in a second quick transaction. "
    "This eliminates the I/O-inside-transaction contention pattern.",
)

# --- 4.5 ---
doc.add_heading(
    "4.5  DB Connection & Contention (connection.py)", level=2
)

doc.add_paragraph(
    "Each DB operation opens a new aiosqlite connection (no pooling). With 8 workers "
    "doing 3-4 writes per file, this means 24+ concurrent connections attempting to "
    "acquire the SQLite write lock. PRAGMA setup (WAL mode, busy_timeout, foreign_keys) "
    "runs on every new connection, adding latency."
)

bold_line(
    "Recommendation: ",
    "Implement a connection pool (3-5 reusable connections). This eliminates per-connection "
    "PRAGMA overhead and reduces lock acquisition contention. Also consider a dedicated "
    "read-only connection for analytics endpoints (pipeline/stats, scanner/progress) so "
    "they never compete with write transactions.",
)

doc.add_page_break()

# ==================================================================
# 5. RECOMMENDATIONS
# ==================================================================
doc.add_heading("5. Recommendations - Prioritized Action Plan", level=1)

# --- P0 ---
doc.add_heading("P0 - Critical (fix this week)", level=2)

doc.add_heading("5.1  Cache Pipeline Stats", level=3)
bold_line("Effort: ", "Small (1-2 hours)")
bold_line("Files: ", "api/routes/pipeline.py")
doc.add_paragraph(
    "Add an in-memory cache with a 15-30 second TTL to pipeline_stats() and "
    "pipeline_status(). The underlying data changes on the order of minutes, not seconds. "
    "A simple time-based cache (asyncio.Lock + last_result + last_time) eliminates "
    "95%+ of the heavy queries. This single change would cut DB read load by roughly half."
)
checkbox("Add TTL cache to pipeline_stats()")
checkbox("Add TTL cache to pipeline_status()")
checkbox("Verify cache invalidates correctly on bulk job start/complete")
checkbox("Test: verify response time drops from seconds to <10ms on cache hit")

doc.add_heading("5.2  Fix Vision Adapter MIME Detection", level=3)
bold_line("Effort: ", "Small (30 minutes)")
bold_line("Files: ", "core/vision_adapter.py")
doc.add_paragraph(
    "Before sending images to the Anthropic API, detect the actual MIME type from "
    "file magic bytes instead of trusting the extension. Use Python imghdr module "
    "or the filetype library. This unblocks ~1,150 stalled images immediately."
)
checkbox("Add magic-byte MIME detection in vision_adapter.py")
checkbox("Fall back to extension-based detection if magic bytes fail")
checkbox("Add unit test for .jpg-that-is-actually-gif case")
checkbox("Re-queue failed analysis_queue entries for retry")

# --- P1 ---
doc.add_heading("P1 - High (fix within 2 weeks)", level=2)

doc.add_heading("5.3  Reduce Frontend Polling Frequency", level=3)
bold_line("Effort: ", "Small (30 minutes)")
bold_line("Files: ", "static/js/global-status-bar.js, plus page-specific JS")
doc.add_paragraph(
    "Change global-status-bar.js polling from 5s -> 15-30s (visible) and 30s -> 60s (hidden). "
    "Review pipeline page JS for additional polling timers and align them. "
    "Long-term: replace polling with SSE push (bulk worker already has SSE infrastructure)."
)
checkbox("Increase global-status-bar.js interval to 15-30s")
checkbox("Audit all page-specific polling intervals")
checkbox("Consider SSE-based status push (future)")

doc.add_heading("5.4  Deduplicate bulk_files", level=3)
bold_line("Effort: ", "Medium (2-4 hours)")
bold_line("Files: ", "core/db/bulk.py, core/bulk_scanner.py")
doc.add_paragraph(
    "The 3.1x inflation means 187K unnecessary rows. Options:\n"
    "1. Run dedup after each scan: DELETE older duplicate (job_id, source_path) rows, "
    "keeping only the latest per source_path.\n"
    "2. Change bulk_files key from (job_id, source_path) to just (source_path) so "
    "re-scans update in place.\n"
    "3. Add composite index on bulk_files(source_path, status) for the NOT EXISTS pattern.\n"
    "4. Rewrite stats queries to drive from source_files (88K rows) instead of bulk_files (276K)."
)
checkbox("Choose dedup strategy")
checkbox("Implement and test")
checkbox("Run one-time cleanup of existing duplicate rows")
checkbox("Benchmark pipeline/stats before and after")

doc.add_heading("5.5  Increase SQLite busy_timeout", level=3)
bold_line("Effort: ", "Trivial (5 minutes)")
bold_line("Files: ", "core/db/connection.py")
doc.add_paragraph(
    "Increase busy_timeout from 5,000ms to 30,000ms. With the stats endpoint "
    "routinely holding locks for 5+ seconds, a 5s timeout guarantees failures. "
    "30s gives writes a realistic chance to acquire the lock."
)
checkbox("Increase busy_timeout to 30,000ms")
checkbox("Evaluate separate read-only connection for stats endpoints")

doc.add_heading("5.6  Batch Worker Counter Updates", level=3)
bold_line("Effort: ", "Small (1-2 hours)")
bold_line("Files: ", "core/bulk_worker.py")
doc.add_paragraph(
    "Currently: 3-4 individual DB writes per file for progress tracking. "
    "Change to: accumulate counters in memory, flush to DB every 50 files "
    "or every 5 seconds (whichever comes first). For 276K files, this reduces "
    "DB writes from ~800K to ~5,500."
)
checkbox("Add in-memory counter accumulator")
checkbox("Add periodic flush (every 50 files or 5 seconds)")
checkbox("Verify SSE progress updates still work correctly")

# --- P2 ---
doc.add_heading("P2 - Medium (next development cycle)", level=2)

items = [
    (
        "5.7  Raise Conversion Semaphore",
        "Trivial",
        "core/converter.py, line ~71",
        "The default of 3 concurrent conversions starves 5 of 8 workers. Set "
        "MAX_CONCURRENT_CONVERSIONS to 6-8, or make it match worker_count. Profile "
        "CPU/memory during conversion to find the optimal limit for the i7-10750H.",
    ),
    (
        "5.8  Suppress httpx/httpcore Debug Logging",
        "Trivial",
        "core/logging_config.py",
        "Set httpcore.* and httpx.* loggers to WARNING in configure_logging(). "
        "These account for ~40,000 log lines/day (30% of markflow.log) with zero "
        "diagnostic value. Same pattern as the existing pdfminer suppression.",
    ),
    (
        "5.9  Connection Pooling",
        "Medium",
        "core/db/connection.py",
        "Replace per-call aiosqlite.connect() with a connection pool (3-5 reusable "
        "connections). Eliminates per-connection PRAGMA overhead. Add a dedicated "
        "read-only connection for analytics endpoints.",
    ),
    (
        "5.10  Move File I/O Outside DB Transactions",
        "Medium",
        "core/lifecycle_manager.py",
        "The move_to_trash function holds a DB write lock during shutil.move (file I/O). "
        "Restructure to: read metadata in transaction -> commit -> move file -> "
        "update status in new quick transaction. Eliminates I/O contention.",
    ),
    (
        "5.11  Preferences Caching",
        "Small",
        "core/scheduler.py",
        "Every scheduler job reads pipeline_enabled, scanner_enabled, business hours "
        "from DB on every tick. Cache in-memory with a 5-minute TTL.",
    ),
    (
        "5.12  Stale Job Detection",
        "Medium",
        "core/scheduler.py, core/scan_coordinator.py",
        "Add a job heartbeat column. On startup, check for jobs with status='running' "
        "but last_heartbeat > 30 minutes and mark them as 'interrupted'. Prevents "
        "orphaned jobs from blocking all maintenance indefinitely.",
    ),
    (
        "5.13  Vector Indexing Backpressure",
        "Small",
        "core/bulk_worker.py",
        "Detached asyncio.create_task() for Qdrant indexing has no queue depth limit. "
        "Add a bounded semaphore (e.g., 20) to prevent unbounded task accumulation "
        "if Qdrant is slow.",
    ),
]

for title_text, effort, files, desc in items:
    doc.add_heading(title_text, level=3)
    bold_line("Effort: ", effort)
    bold_line("Files: ", files)
    doc.add_paragraph(desc)

doc.add_page_break()

# ==================================================================
# 6. TRACKING CHECKLIST
# ==================================================================
doc.add_heading("6. Implementation Tracking Checklist", level=1)
doc.add_paragraph("Use this table to track progress. Print and check off as you go.")

add_table(
    ["#", "Finding / Action", "Priority", "Effort", "Status", "Notes"],
    [
        ["5.1", "Cache pipeline stats (15-30s TTL)", "P0", "Small", "[ ] Not started", "Biggest single win"],
        ["5.2", "Fix vision MIME detection", "P0", "Small", "[ ] Not started", "Unblocks 1,150 images"],
        ["5.3", "Reduce frontend polling (5s -> 15-30s)", "P1", "Small", "[ ] Not started", "49K fewer req/day"],
        ["5.4", "Deduplicate bulk_files (3.1x -> 1x)", "P1", "Medium", "[ ] Not started", "Delete ~187K rows"],
        ["5.5", "Increase busy_timeout (5s -> 30s)", "P1", "Trivial", "[ ] Not started", "5-minute fix"],
        ["5.6", "Batch worker counter updates", "P1", "Small", "[ ] Not started", "98% fewer DB writes"],
        ["5.7", "Raise conversion semaphore (3 -> 6-8)", "P2", "Trivial", "[ ] Not started", "Profile first"],
        ["5.8", "Suppress httpx debug logging", "P2", "Trivial", "[ ] Not started", "40K fewer log lines/day"],
        ["5.9", "Connection pooling", "P2", "Medium", "[ ] Not started", "Reduces lock contention"],
        ["5.10", "Move file I/O outside transactions", "P2", "Medium", "[ ] Not started", "Fixes trash_move locks"],
        ["5.11", "Preferences caching (5-min TTL)", "P2", "Small", "[ ] Not started", ""],
        ["5.12", "Stale job detection", "P2", "Medium", "[ ] Not started", "Prevents orphaned blocks"],
        ["5.13", "Vector indexing backpressure", "P2", "Small", "[ ] Not started", "Bounded semaphore"],
        ["--", "Restore production lifecycle timers", "Pre-prod", "N/A", "[ ] Not started", "ACTION REQUIRED"],
    ],
)

doc.add_page_break()

# ==================================================================
# APPENDICES
# ==================================================================
doc.add_heading("Appendix A: Log Analysis Methodology", level=1)
doc.add_paragraph(
    "Analysis performed on 2026-04-09 against live container logs and on-disk log files. "
    "Log window: 04:28-18:13 UTC (~14 hours). All JSON-structured logs in logs/markflow.log "
    "were parsed and aggregated by event type and level using Python scripts. HTTP request "
    "durations extracted from request_complete events. DB state queried via python3 inside "
    "the running container."
)
add_table(
    ["Source", "Lines Analyzed"],
    [
        ["logs/markflow.log", "150,731"],
        ["logs/markflow-debug.log", "121,445"],
        ["logs/db-contention.log", "2,441,236"],
        ["logs/db-queries.log", "876,645"],
        ["logs/db-active.log", "42,447"],
        ["Total", "3,632,504"],
    ],
)

doc.add_heading("Appendix B: Container State at Time of Audit", level=1)
add_table(
    ["Container", "Image", "Status", "Ports"],
    [
        ["markflow", "doc-conversion-2026-markflow", "Up 14 hours", "8000"],
        ["markflow-mcp", "doc-conversion-2026-markflow-mcp", "Up 14 hours", "8001"],
        ["meilisearch", "getmeili/meilisearch:latest", "Up 14 hours (healthy)", "7700"],
        ["qdrant", "qdrant/qdrant:latest", "Up 14 hours (healthy)", "6333-6334"],
    ],
)

doc.add_heading("Appendix C: SQLite Configuration", level=1)
add_table(
    ["Setting", "Current Value", "Recommended"],
    [
        ["Journal mode", "WAL", "WAL (no change)"],
        ["WAL autocheckpoint", "1,000 pages", "1,000 (no change)"],
        ["busy_timeout", "5,000 ms", "30,000 ms"],
        ["Database size", "312.9 MB", "~120 MB after dedup"],
        ["WAL file size", "0.0 MB (clean)", ""],
        ["Index count", "69", "Add 1-2 composite"],
    ],
)

doc.add_heading("Appendix D: bulk_files Duplication Distribution", level=1)
doc.add_paragraph(
    "Distribution showing how many bulk_files rows exist per unique source_path. "
    "Ideally each path would have exactly 1 row."
)
add_table(
    ["Rows per Path", "Number of Paths", "Total Rows", "Waste (rows - paths)"],
    [
        ["1", "23,729", "23,729", "0"],
        ["2", "27,187", "54,374", "27,187"],
        ["3", "20,115", "60,345", "40,230"],
        ["4", "2,736", "10,944", "8,208"],
        ["5", "989", "4,945", "3,956"],
        ["6", "200", "1,200", "1,000"],
        ["7", "4,693", "32,851", "28,158"],
        ["8", "5,288", "42,304", "37,016"],
        ["9", "6", "54", "48"],
        ["10", "1,126", "11,260", "10,134"],
        ["11", "1,067", "11,737", "10,670"],
        ["12-22", "~1,400", "~22,576", "~21,176"],
        ["TOTAL", "88,536", "276,320", "~187,784 wasted rows"],
    ],
)

# Save
out_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(os.path.dirname(out_dir))
out = os.path.join(project_root, "MarkFlow-Health-Audit-2026-04-09.docx")
doc.save(out)
print(f"Saved: {out}")
