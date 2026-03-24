"""
DOCX/DOC format handler — bidirectional conversion via python-docx.

DocxHandler.ingest(file_path) → DocumentModel:
  Extracts headings, paragraphs, tables (nested), inline images, footnotes.
  .doc files are first converted to .docx via LibreOffice headless.

DocxHandler.extract_styles(file_path) → dict:
  Per-element font/size/spacing/color, table structure, document-level page settings.

DocxHandler.export(model, output_path, sidecar=None):
  Tier 1: structure always. Tier 2: styles from sidecar. Tier 3: patch against original.
"""

import logging
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from formats.base import FormatHandler, register_handler
from core.document_model import (
    DocumentModel,
    DocumentMetadata,
    Element,
    ElementType,
    ImageData,
    compute_content_hash,
)
from core.image_handler import extract_image

log = logging.getLogger(__name__)

# DOCX paragraph style name → (ElementType, level)
_HEADING_STYLES: dict[str, tuple[ElementType, int]] = {
    "heading 1": (ElementType.HEADING, 1),
    "heading 2": (ElementType.HEADING, 2),
    "heading 3": (ElementType.HEADING, 3),
    "heading 4": (ElementType.HEADING, 4),
    "heading 5": (ElementType.HEADING, 5),
    "heading 6": (ElementType.HEADING, 6),
    "title": (ElementType.HEADING, 1),
    "subtitle": (ElementType.HEADING, 2),
}

_BLOCKQUOTE_STYLES = {"block text", "quote", "quotations", "intense quote"}
_CODE_STYLES = {"code", "html code", "code block"}


# ── Utility helpers ───────────────────────────────────────────────────────────

def _para_text_with_formatting(para) -> str:
    """
    Build a plain-text representation of a paragraph with basic Markdown
    inline formatting (bold, italic).
    """
    parts = []
    for run in para.runs:
        text = run.text or ""
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def _style_name(para) -> str:
    """Return the lowercase style name for a paragraph (or '' if none)."""
    try:
        return (para.style.name or "").lower().strip()
    except Exception:
        return ""


def _is_list_para(para) -> bool:
    """Return True if the paragraph is a list item (has numPr or list-style name)."""
    try:
        pPr = para._element.pPr
        if pPr is not None and pPr.numPr is not None:
            return True
        sname = _style_name(para)
        return any(x in sname for x in ("list bullet", "list number", "list paragraph"))
    except Exception:
        return False


def _list_level(para) -> int:
    """Return the 0-based list nesting level."""
    try:
        pPr = para._element.pPr
        if pPr is None or pPr.numPr is None:
            return 0
        ilvl = pPr.numPr.ilvl
        return int(ilvl.val) if ilvl is not None else 0
    except Exception:
        return 0


def _is_ordered_list(para) -> bool:
    """Heuristic: detect ordered list by checking style name."""
    sname = _style_name(para)
    return "list number" in sname or "numbered" in sname


def _emu_to_pt(emu: int | None) -> float | None:
    """Convert EMU (English Metric Units) to points."""
    if emu is None:
        return None
    return round(emu / 12700, 2)


def _pt_to_emu(pt: float) -> int:
    return int(pt * 12700)


def _plain_text_hash(text: str) -> str:
    """
    Strip inline markdown markers (**bold**, *italic*, `code`) and return a
    content hash for sidecar lookup.

    Sidecar keys are generated from para.text (plain) during DOCX ingest, but
    element content may include markdown markers after re-ingest of Markdown.
    Using plain text ensures the hashes match.
    """
    plain = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", str(text), flags=re.DOTALL)
    plain = re.sub(r"`(.+?)`", r"\1", plain, flags=re.DOTALL)
    return compute_content_hash(plain)


def _add_inline_runs(para, text: str) -> None:
    """
    Parse inline markdown markers in *text* and append properly formatted
    Word runs to *para*.

    Handles:  ***bold+italic***  **bold**  *italic*  `code`  and plain text.
    """
    # Split on formatting tokens; keep delimiters via the capture group
    parts = re.split(
        r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)",
        text,
        flags=re.DOTALL,
    )
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            run = para.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            para.add_run(part)


# ── Image extraction ──────────────────────────────────────────────────────────

def _extract_para_images(para, doc) -> list[tuple[str, bytes, dict[str, Any]]]:
    """
    Extract inline images from a paragraph's runs.

    Returns list of (hash_filename, png_data, metadata).
    """
    from docx.oxml.ns import qn

    results = []
    for run in para.runs:
        for drawing in run._element.iter(qn("a:blip")):
            r_embed = drawing.get(qn("r:embed"))
            if r_embed and r_embed in doc.part.related_parts:
                image_part = doc.part.related_parts[r_embed]
                content_type = image_part.content_type  # e.g. "image/png"
                ext = content_type.split("/")[-1] if "/" in content_type else "png"
                # Normalize some MIME subtypes
                ext = {"jpeg": "jpg", "tiff": "tif"}.get(ext, ext)
                try:
                    fname, png_data, meta = extract_image(image_part.blob, ext)
                    results.append((fname, png_data, meta))
                except Exception as exc:
                    log.warning("docx.image_extract_failed", error=str(exc))
    return results


# ── .doc → .docx conversion via LibreOffice ───────────────────────────────────

def _doc_to_docx(doc_path: Path) -> Path:
    """
    Convert a .doc file to .docx using LibreOffice headless.
    Returns path to the new .docx file (in a temp directory).
    Raises RuntimeError if LibreOffice is not available or conversion fails.
    """
    for binary in ("libreoffice", "soffice"):
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    [binary, "--headless", "--convert-to", "docx",
                     "--outdir", tmpdir, str(doc_path)],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
                if result.returncode == 0:
                    out_path = Path(tmpdir) / (doc_path.stem + ".docx")
                    if out_path.exists():
                        # Copy to a stable temp file outside tmpdir
                        import shutil
                        stable = Path(tempfile.mktemp(suffix=".docx"))
                        shutil.copy2(out_path, stable)
                        return stable
        except FileNotFoundError:
            continue
    raise RuntimeError(
        f"Cannot convert {doc_path.name}: LibreOffice not found. "
        "Install libreoffice-headless."
    )


# ── Table extraction ──────────────────────────────────────────────────────────

def _extract_table(table, doc, model: DocumentModel, depth: int = 0) -> Element:
    """
    Recursively extract a python-docx Table into a TABLE Element.
    Nested tables are represented as cell text "[nested table]".
    """
    rows: list[list[str]] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            # Check for nested table
            cell_tables = cell.tables
            if cell_tables and depth < 2:
                # Represent nested table as a special placeholder
                nested_texts = []
                for nested in cell_tables:
                    nested_elem = _extract_table(nested, doc, model, depth + 1)
                    # Flatten the nested table to a text representation
                    nested_rows = nested_elem.content if isinstance(nested_elem.content, list) else []
                    nested_texts.append(" / ".join(
                        " | ".join(str(c) for c in row) for row in nested_rows
                    ))
                cells.append(" [table: " + " ; ".join(nested_texts) + "]")
            else:
                # Get cell text (all paragraphs joined)
                cell_text = "\n".join(
                    _para_text_with_formatting(p) for p in cell.paragraphs
                ).strip()
                cells.append(cell_text)
        rows.append(cells)

    return Element(
        type=ElementType.TABLE,
        content=rows,
    )


# ── Footnote extraction ───────────────────────────────────────────────────────

def _extract_footnotes(doc) -> list[tuple[str, str]]:
    """
    Extract footnote definitions from a DOCX document.
    Returns list of (id, text) tuples.
    """
    from docx.oxml.ns import qn

    footnotes = []
    try:
        fn_part = doc.part.footnotes_part
        if fn_part is None:
            return []
        for fn in fn_part._element.findall(qn("w:footnote")):
            fn_id = fn.get(qn("w:id"), "")
            # Skip separator footnotes (id -1 and 0)
            if fn_id in ("-1", "0"):
                continue
            paragraphs = fn.findall(qn("w:p"))
            text = " ".join(
                "".join(r.text or "" for r in p.iter(qn("w:t")))
                for p in paragraphs
            ).strip()
            if text:
                footnotes.append((fn_id, text))
    except Exception as exc:
        log.debug("docx.footnote_extract_skip", reason=str(exc))
    return footnotes


# ── Main handler ──────────────────────────────────────────────────────────────

@register_handler
class DocxHandler(FormatHandler):
    """DOCX/DOC format handler using python-docx."""

    EXTENSIONS = ["docx", "doc"]

    # ── Ingest ────────────────────────────────────────────────────────────────

    def ingest(self, file_path: Path) -> DocumentModel:
        """Read a .docx (or .doc) file and return a DocumentModel."""
        import docx

        file_path = Path(file_path)
        _tmp_docx: Path | None = None

        if file_path.suffix.lower() == ".doc":
            # Convert .doc → .docx first
            file_path = _doc_to_docx(file_path)
            _tmp_docx = file_path

        try:
            doc = docx.Document(str(file_path))
        finally:
            if _tmp_docx and _tmp_docx.exists():
                _tmp_docx.unlink(missing_ok=True)

        model = DocumentModel()

        # ── Core properties (title, author, subject) ──────────────────────
        cp = doc.core_properties
        model.metadata = DocumentMetadata(
            source_file=Path(file_path).name,
            source_format="docx",
            title=cp.title or "",
            author=cp.author or "",
            subject=cp.subject or "",
        )

        # ── Walk document body in element order ───────────────────────────
        from docx.oxml.ns import qn

        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                self._process_paragraph(para, doc, model)

            elif tag == "tbl":
                from docx.table import Table
                table = Table(child, doc)
                elem = _extract_table(table, doc, model)
                if isinstance(elem.content, list) and elem.content:
                    model.add_element(elem)

            elif tag == "sdt":
                # Structured document tag — try to extract inner text
                for p_elem in child.iter(qn("w:p")):
                    from docx.text.paragraph import Paragraph
                    para = Paragraph(p_elem, doc)
                    self._process_paragraph(para, doc, model)

        # ── Footnotes ─────────────────────────────────────────────────────
        for fn_id, fn_text in _extract_footnotes(doc):
            model.add_element(Element(
                type=ElementType.FOOTNOTE,
                content=fn_text,
                attributes={"id": fn_id},
            ))

        return model

    def _process_paragraph(self, para, doc, model: DocumentModel) -> None:
        """Process a single paragraph and add element(s) to model."""
        style = _style_name(para)

        # ── Heading ───────────────────────────────────────────────────────
        if style in _HEADING_STYLES:
            elem_type, level = _HEADING_STYLES[style]
            text = para.text.strip()
            if text:
                model.add_element(Element(
                    type=elem_type,
                    content=text,
                    level=level,
                    attributes={"style_name": para.style.name},
                ))
            return

        # ── Code block ────────────────────────────────────────────────────
        if style in _CODE_STYLES:
            text = para.text
            model.add_element(Element(
                type=ElementType.CODE_BLOCK,
                content=text,
            ))
            return

        # ── Blockquote ────────────────────────────────────────────────────
        if style in _BLOCKQUOTE_STYLES:
            text = _para_text_with_formatting(para).strip()
            if text:
                model.add_element(Element(
                    type=ElementType.BLOCKQUOTE,
                    content=text,
                ))
            return

        # ── Images in paragraph ───────────────────────────────────────────
        images_in_para = _extract_para_images(para, doc)
        for fname, png_data, meta in images_in_para:
            model.images[fname] = ImageData(
                data=png_data,
                original_format=meta.get("original_format", "png"),
                width=meta.get("width"),
                height=meta.get("height"),
                alt_text=para.text.strip() or "",
            )
            model.add_element(Element(
                type=ElementType.IMAGE,
                content="",
                attributes={
                    "src": f"assets/{fname}",
                    "alt": para.text.strip() or "image",
                    "width": meta.get("width"),
                    "height": meta.get("height"),
                },
            ))

        # If paragraph had images, skip adding paragraph text element
        if images_in_para:
            return

        # ── List item ─────────────────────────────────────────────────────
        if _is_list_para(para):
            text = _para_text_with_formatting(para).strip()
            if text:
                ordered = _is_ordered_list(para)
                level = _list_level(para)
                model.add_element(Element(
                    type=ElementType.LIST_ITEM,
                    content=text,
                    level=level,
                    attributes={"ordered": ordered},
                ))
            return

        # ── Page break ────────────────────────────────────────────────────
        from docx.oxml.ns import qn
        if para._element.find(qn("w:lastRenderedPageBreak")) is not None or (
            para.runs and any(
                r._element.find(qn("w:br")) is not None
                and r._element.find(qn("w:br")).get(qn("w:type")) == "page"
                for r in para.runs
            )
        ):
            model.add_element(Element(type=ElementType.PAGE_BREAK, content=""))
            return

        # ── Paragraph ─────────────────────────────────────────────────────
        text = _para_text_with_formatting(para).strip()
        if text:
            model.add_element(Element(
                type=ElementType.PARAGRAPH,
                content=text,
                attributes={"style_name": para.style.name if para.style else ""},
            ))

    # ── Extract styles ────────────────────────────────────────────────────────

    def extract_styles(self, file_path: Path) -> dict[str, Any]:
        """
        Extract per-element and document-level style data.

        Returns dict keyed by content hash, plus a 'document_level' key.
        """
        import docx

        file_path = Path(file_path)
        _tmp_docx: Path | None = None

        if file_path.suffix.lower() == ".doc":
            file_path = _doc_to_docx(file_path)
            _tmp_docx = file_path

        try:
            doc = docx.Document(str(file_path))
        finally:
            if _tmp_docx and _tmp_docx.exists():
                _tmp_docx.unlink(missing_ok=True)

        style_data: dict[str, Any] = {}

        # ── Document-level settings ───────────────────────────────────────
        section = doc.sections[0] if doc.sections else None
        doc_level: dict[str, Any] = {}
        if section:
            doc_level["page_width_pt"] = _emu_to_pt(section.page_width)
            doc_level["page_height_pt"] = _emu_to_pt(section.page_height)
            doc_level["margin_top_pt"] = _emu_to_pt(section.top_margin)
            doc_level["margin_bottom_pt"] = _emu_to_pt(section.bottom_margin)
            doc_level["margin_left_pt"] = _emu_to_pt(section.left_margin)
            doc_level["margin_right_pt"] = _emu_to_pt(section.right_margin)

        # Default font from document styles
        try:
            normal_style = doc.styles["Normal"]
            font = normal_style.font
            doc_level["default_font"] = font.name or ""
            doc_level["default_size_pt"] = (
                font.size.pt if font.size else None
            )
        except (KeyError, AttributeError):
            pass

        style_data["document_level"] = doc_level

        # ── Per-paragraph styles ──────────────────────────────────────────
        from docx.oxml.ns import qn

        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue
                h = compute_content_hash(text)
                entry = self._para_style_entry(para)
                if entry:
                    style_data[h] = entry

            elif tag == "tbl":
                from docx.table import Table
                table = Table(child, doc)
                self._table_styles(table, style_data)

        return style_data

    def _para_style_entry(self, para) -> dict[str, Any] | None:
        """Build a style dict for a single paragraph."""
        try:
            font = para.runs[0].font if para.runs else None
            pf = para.paragraph_format
            entry: dict[str, Any] = {
                "type": "paragraph",
                "style_name": para.style.name if para.style else "",
            }
            if font:
                entry["font_family"] = font.name or ""
                entry["font_size_pt"] = font.size.pt if font.size else None
                entry["bold"] = font.bold or False
                entry["italic"] = font.italic or False
                entry["underline"] = font.underline or False
                if font.color and font.color.rgb:
                    entry["color"] = str(font.color.rgb)
            if pf:
                entry["alignment"] = str(pf.alignment) if pf.alignment else "left"
                entry["space_before_pt"] = (
                    pf.space_before.pt if pf.space_before else None
                )
                entry["space_after_pt"] = (
                    pf.space_after.pt if pf.space_after else None
                )
                entry["line_spacing"] = (
                    float(pf.line_spacing) if pf.line_spacing else None
                )
            return entry
        except Exception:
            return None

    def _table_styles(self, table, style_data: dict[str, Any]) -> None:
        """Extract and store style info for a table."""
        try:
            rows_text = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            h = compute_content_hash(rows_text)
            entry: dict[str, Any] = {"type": "table"}
            # Column widths in points
            col_widths = []
            for col in table.columns:
                try:
                    col_widths.append(_emu_to_pt(col.width))
                except Exception:
                    col_widths.append(None)
            if col_widths:
                entry["column_widths_pt"] = col_widths
            # Table style name
            try:
                entry["table_style"] = table.style.name if table.style else ""
            except Exception:
                pass
            style_data[h] = entry
        except Exception as exc:
            log.debug("docx.table_style_skip", reason=str(exc))

    # ── Export ────────────────────────────────────────────────────────────────

    def export(
        self,
        model: DocumentModel,
        output_path: Path,
        sidecar: dict[str, Any] | None = None,
        original_path: Path | None = None,
    ) -> None:
        """
        Write a DocumentModel to a .docx file.

        Tier 1 (always): headings, paragraphs, tables, images, lists.
        Tier 2 (if sidecar): apply font/spacing/alignment from sidecar.
        Tier 3 (if original_path): use original as template base, then rebuild.
        """
        import docx
        from docx.shared import Pt

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # ── Tier 3: try original template ─────────────────────────────────
        if original_path and Path(original_path).exists():
            doc = self._patch_from_original(model, Path(original_path), sidecar, output_path)
            if doc is not None:
                doc.save(str(output_path))
                return

        # ── Tier 1/2: build from scratch ──────────────────────────────────
        doc = docx.Document()

        # Apply document-level settings from sidecar (Tier 2)
        if sidecar:
            dl = sidecar.get("document_level", {})
            section = doc.sections[0]
            if dl.get("margin_top_pt"):
                section.top_margin = Pt(dl["margin_top_pt"])
            if dl.get("margin_bottom_pt"):
                section.bottom_margin = Pt(dl["margin_bottom_pt"])
            if dl.get("margin_left_pt"):
                section.left_margin = Pt(dl["margin_left_pt"])
            if dl.get("margin_right_pt"):
                section.right_margin = Pt(dl["margin_right_pt"])

        for elem in model.elements:
            self._export_element(doc, elem, model, sidecar, output_path)

        doc.save(str(output_path))

    def _export_element(
        self,
        doc,
        elem: Element,
        model: DocumentModel,
        sidecar: dict[str, Any] | None,
        output_path: Path,
    ) -> None:
        """Write a single Element to the python-docx Document."""
        import docx as _docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        t = elem.type

        if t == ElementType.HEADING:
            level = min(max(elem.level or 1, 1), 9)
            doc.add_heading(str(elem.content), level=level)

        elif t == ElementType.PARAGRAPH:
            p = doc.add_paragraph()
            _add_inline_runs(p, str(elem.content))
            self._apply_sidecar_style(p, str(elem.content), sidecar)

        elif t == ElementType.CODE_BLOCK:
            p = doc.add_paragraph(str(elem.content), style="Normal")
            p.runs[0].font.name = "Courier New"

        elif t == ElementType.BLOCKQUOTE:
            p = doc.add_paragraph(str(elem.content))
            fmt = p.paragraph_format
            from docx.shared import Inches
            fmt.left_indent = Inches(0.5)
            fmt.right_indent = Inches(0.5)

        elif t == ElementType.HORIZONTAL_RULE:
            # Approximate with a bottom-border paragraph
            doc.add_paragraph("─" * 40)

        elif t == ElementType.PAGE_BREAK:
            doc.add_page_break()

        elif t == ElementType.LIST_ITEM:
            ordered = elem.attributes.get("ordered", False)
            style_name = "List Number" if ordered else "List Bullet"
            try:
                doc.add_paragraph(str(elem.content), style=style_name)
            except KeyError:
                doc.add_paragraph(("1. " if ordered else "• ") + str(elem.content))

        elif t == ElementType.LIST:
            if elem.children:
                for child in elem.children:
                    self._export_element(doc, child, model, sidecar, output_path)

        elif t == ElementType.TABLE:
            rows = elem.content if isinstance(elem.content, list) else []
            if not rows:
                return
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < cols:
                        table.cell(r_idx, c_idx).text = str(cell_text)

            # Apply column widths from sidecar (Tier 2)
            if sidecar:
                elements_map = sidecar.get("elements", {})
                entry = elements_map.get(elem.content_hash)
                if entry and entry.get("type") == "table":
                    col_widths = entry.get("column_widths_pt", [])
                    for c_idx, col_pt in enumerate(col_widths):
                        if c_idx < cols and col_pt:
                            try:
                                from docx.shared import Pt as _Pt
                                for trow in table.rows:
                                    trow.cells[c_idx].width = _Pt(col_pt)
                            except Exception:
                                pass

        elif t == ElementType.IMAGE:
            src = elem.attributes.get("src", "")
            # Try to find the image in the model or relative to output_path
            img_name = Path(src).name
            if img_name in model.images:
                img_data = model.images[img_name].data
                try:
                    width = elem.attributes.get("width")
                    pic_kwargs: dict = {}
                    if width:
                        pic_kwargs["width"] = Inches(min(width / 96, 6.0))
                    doc.add_picture(io.BytesIO(img_data), **pic_kwargs)
                except Exception as exc:
                    log.warning("docx.export_image_failed", src=src, error=str(exc))
                    doc.add_paragraph(f"[Image: {elem.attributes.get('alt', src)}]")
            else:
                # Image file may be on disk
                assets_dir = output_path.parent / "assets"
                img_path = assets_dir / img_name
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path))
                    except Exception:
                        doc.add_paragraph(f"[Image: {elem.attributes.get('alt', src)}]")
                else:
                    doc.add_paragraph(f"[Image: {elem.attributes.get('alt', src)}]")

        elif t == ElementType.FOOTNOTE:
            # Represent as a paragraph note (true footnotes require XML manipulation)
            doc.add_paragraph(f"^[{elem.attributes.get('id', '')}]: {elem.content}")

        elif t == ElementType.RAW_HTML:
            # Strip HTML tags for plain text
            import re
            text = re.sub(r"<[^>]+>", "", str(elem.content))
            if text.strip():
                doc.add_paragraph(text)

    def _apply_sidecar_style(self, para, content: str, sidecar: dict | None) -> None:
        """
        Apply Tier 2 style from sidecar to a paragraph (best-effort).

        Looks up the sidecar by plain-text hash (markdown markers stripped) so
        that the lookup works regardless of whether inline formatting markers
        are present in the element content.
        """
        if not sidecar:
            return
        elements_map = sidecar.get("elements", {})
        # Strip markdown markers before hashing (sidecar keys use plain text)
        entry = elements_map.get(_plain_text_hash(content))
        if not entry:
            # Fallback: try direct hash (content without markers)
            entry = elements_map.get(compute_content_hash(content))
        if not entry:
            return
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Apply run-level formatting to ALL runs
            for run in para.runs:
                if entry.get("font_family"):
                    run.font.name = entry["font_family"]
                if entry.get("font_size_pt"):
                    run.font.size = Pt(entry["font_size_pt"])
                if entry.get("bold") is True:
                    run.font.bold = True
                if entry.get("italic") is True:
                    run.font.italic = True
                if entry.get("color"):
                    hex_color = str(entry["color"]).lstrip("#")
                    if len(hex_color) == 6:
                        r, g, b = (
                            int(hex_color[0:2], 16),
                            int(hex_color[2:4], 16),
                            int(hex_color[4:6], 16),
                        )
                        run.font.color.rgb = RGBColor(r, g, b)

            # Apply paragraph-level formatting
            fmt = para.paragraph_format
            if entry.get("space_before_pt") is not None:
                fmt.space_before = Pt(entry["space_before_pt"])
            if entry.get("space_after_pt") is not None:
                fmt.space_after = Pt(entry["space_after_pt"])
            if entry.get("line_spacing") is not None:
                try:
                    fmt.line_spacing = float(entry["line_spacing"])
                except (TypeError, ValueError):
                    pass

            align_str = str(entry.get("alignment") or "")
            _align_map = {
                "WD_ALIGN_PARAGRAPH.LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                "WD_ALIGN_PARAGRAPH.CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                "WD_ALIGN_PARAGRAPH.RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                "WD_ALIGN_PARAGRAPH.JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if align_str in _align_map:
                fmt.alignment = _align_map[align_str]

        except Exception as exc:
            log.debug("docx.apply_sidecar_style_failed", content=content[:40], reason=str(exc))

    # ── Tier 3: patch from original ───────────────────────────────────────────

    def _patch_from_original(
        self,
        model: DocumentModel,
        original_path: Path,
        sidecar: dict[str, Any] | None,
        output_path: Path,
    ):
        """
        Tier 3: Use the original DOCX as a template base.

        Computes match ratio between original and model content hashes.
        If ≥ 80% of model elements match the original, opens the original,
        clears its body, and rebuilds content from the model using Tier 2
        styling — this preserves the original's registered styles, themes,
        and document settings.

        Returns a Document on success, or None to fall back to Tier 2.
        """
        import docx as _docx

        try:
            orig_doc = _docx.Document(str(original_path))
        except Exception as exc:
            log.warning("docx.tier3_open_failed", error=str(exc))
            return None

        # Build set of content hashes from original paragraphs
        orig_hashes: set[str] = set()
        for child in orig_doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                from docx.text.paragraph import Paragraph as _Para
                para = _Para(child, orig_doc)
                text = para.text.strip()
                if text:
                    orig_hashes.add(compute_content_hash(text))

        # Build list of plain-text hashes from model elements
        model_hashes: list[str] = []
        for elem in model.elements:
            if isinstance(elem.content, str) and elem.content.strip():
                plain = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", elem.content, flags=re.DOTALL)
                plain = re.sub(r"`(.+?)`", r"\1", plain, flags=re.DOTALL)
                model_hashes.append(compute_content_hash(plain))

        if not model_hashes:
            return None

        matching = sum(1 for h in model_hashes if h in orig_hashes)
        match_ratio = matching / len(model_hashes)

        log.info(
            "docx.tier3_match_ratio",
            matching=matching,
            total=len(model_hashes),
            ratio=round(match_ratio, 2),
        )

        if match_ratio < 0.8:
            log.info("docx.tier3_fallback", reason="match_ratio_below_0.8", ratio=round(match_ratio, 2))
            return None

        # Clear body (keep sectPr to preserve page layout)
        from docx.oxml.ns import qn as _qn
        body = orig_doc.element.body
        sect_pr = body.find(_qn("w:sectPr"))
        for child in list(body):
            if child is not sect_pr:
                body.remove(child)

        # Rebuild body from model using Tier 2 styling
        for elem in model.elements:
            self._export_element(orig_doc, elem, model, sidecar, output_path)

        log.info("docx.tier3_complete", file=str(output_path))
        return orig_doc
